import io
import streamlit as st

from sintetizador import gerar_sintese_academica
from reescrever_texto import reescrever_texto

from originalidade import (
    executar_analise_originalidade,
    gerar_relatorio_originalidade,
    construir_prompt_parafrase,
    comparar_antes_depois,
)

from ia_detector import (
    analisar_texto_ia,
    gerar_relatorio_ia,
)


# ============================================================
# CONFIGURAÇÃO
# ============================================================

st.set_page_config(
    page_title="DocScripta AI",
    page_icon="📚",
    layout="wide",
    menu_items={
        "About": "DocScripta AI — Assistente acadêmico inteligente"
    }
)


# ============================================================
# PWA — LINKS PARA MANIFEST E SERVICE WORKER
# ============================================================

pwa_html = """
<link rel="manifest" href="/static/manifest.json">
<meta name="theme-color" content="" id="pwa-theme-color">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
<meta name="apple-mobile-web-app-title" content="DocScripta AI">
<link rel="apple-touch-icon" href="/static/tonnybot-192.png">
<script>
if ('serviceWorker' in navigator) {
  navigator.serviceWorker.register('/static/sw.js');
}
</script>
"""

st.components.v1.html(pwa_html, height=0)


# ============================================================
# TOGGLE DE TEMA — CLARO / ESCURO
# ============================================================

def aplicar_tema():
    tema = st.session_state.get("theme_mode", "dark")

    if tema == "dark":
        cor_fundo = "#0e1117"
        cor_fundo_secundario = "#1a1a2e"
        cor_texto = "#fafafa"
        cor_texto_secundario = "#a0a0b0"
        cor_accent = "#5b6abf"
        cor_card = "#1e1e30"
        cor_borda = "#2a2a40"
        cor_sucesso = "#d1fae5"
        cor_sucesso_texto = "#065f46"
        cor_warn = "#fef3c7"
        cor_warn_texto = "#92400e"
        cor_info = "#dbeafe"
        cor_info_texto = "#1e40af"
        cor_botao = "#5b6abf"
        cor_botao_texto = "#ffffff"
        cor_input_bg = "#1a1a2e"
        cor_input_text = "#fafafa"
        cor_input_borda = "#2a2a40"
        pwa_color = "#0e1117"
    else:
        cor_fundo = "#ffffff"
        cor_fundo_secundario = "#f0f2f6"
        cor_texto = "#1a1a2e"
        cor_texto_secundario = "#555570"
        cor_accent = "#4a59a0"
        cor_card = "#f7f8fa"
        cor_borda = "#e0e0e8"
        cor_sucesso = "#d1fae5"
        cor_sucesso_texto = "#065f46"
        cor_warn = "#fef3c7"
        cor_warn_texto = "#92400e"
        cor_info = "#dbeafe"
        cor_info_texto = "#1e40af"
        cor_botao = "#4a59a0"
        cor_botao_texto = "#ffffff"
        cor_input_bg = "#ffffff"
        cor_input_text = "#1a1a2e"
        cor_input_borda = "#d0d0d8"
        pwa_color = "#ffffff"

    tema_css = f"""    <style>

    /* === ESCONDER STREAMLIT === */
    #MainMenu {{visibility: hidden;}}
    footer {{visibility: hidden;}}
    header {{visibility: hidden;}}

    /* === FUNDO GLOBAL === */
    .stApp {{
        background-color: {cor_fundo} !important;
        color: {cor_texto} !important;
    }}

    /* === SIDEBAR === */
    section[data-testid="stSidebar"] {{
        background-color: {cor_fundo_secundario} !important;
    }}

    /* === TEXTO E TÍTULOS === */
    .stMarkdown, .stText, p, span {{
        color: {cor_texto} !important;
    }}

    h1, h2, h3, h4, h5, h6 {{
        color: {cor_texto} !important;
    }}

    /* === CAPTIONS === */
    .stCaption {{
        color: {cor_texto_secundario} !important;
    }}

    /* === INPUTS / TEXT AREAS === */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {{
        background-color: {cor_input_bg} !important;
        color: {cor_input_text} !important;
        border-color: {cor_input_borda} !important;
    }}

    /* === SELECTBOX / RADIO === */
    .stSelectbox > div > div,
    .stRadio > div {{
        color: {cor_texto} !important;
    }}

    /* === BOTÕES === */
    .stButton > button {{
        background-color: {cor_botao} !important;
        color: {cor_botao_texto} !important;
        border-color: {cor_accent} !important;
        border-radius: 8px !important;
    }}

    .stButton > button:hover {{
        opacity: 0.85 !important;
    }}

    /* === PRIMARY BUTTON === */
    .stButton > button[kind="primary"] {{
        background-color: {cor_accent} !important;
        color: #ffffff !important;
        border-color: {cor_accent} !important;
    }}

    /* === CARDS / EXPANDERS === */
    .stExpander > details {{
        background-color: {cor_card} !important;
        border-color: {cor_borda} !important;
    }}

    .stExpander > details > summary {{
        color: {cor_texto} !important;
    }}

    .stExpander > details > div {{
        color: {cor_texto} !important;
    }}

    /* === METRICS === */
    [data-testid="stMetricValue"] {{
        color: {cor_texto} !important;
    }}

    [data-testid="stMetricLabel"] {{
        color: {cor_texto_secundario} !important;
    }}

    /* === DOWNLOAD BUTTON === */
    .stDownloadButton > button {{
        background-color: {cor_accent} !important;
        color: #ffffff !important;
        border-color: {cor_accent} !important;
    }}

    /* === SLIDER === */
    .stSlider {{
        color: {cor_texto} !important;
    }}

    /* === CHECKBOX === */
    .stCheckbox {{
        color: {cor_texto} !important;
    }}

    /* === FORM === */
    .stForm {{
        background-color: {cor_fundo_secundario} !important;
        border-color: {cor_borda} !important;
    }}

    /* === DIVIDER === */
    .stDivider {{
        border-color: {cor_borda} !important;
    }}

    /* === SUCCESS / WARNING / ERROR / INFO === */
    .stSuccess {{
        background-color: {cor_sucesso} !important;
        color: {cor_sucesso_texto} !important;
    }}

    .stWarning {{
        background-color: {cor_warn} !important;
        color: {cor_warn_texto} !important;
    }}

    .stInfo {{
        background-color: {cor_info} !important;
        color: {cor_info_texto} !important;
    }}

    .stError {{
        background-color: #fde8e8 !important;
        color: #9b1c1c !important;
    }}

    /* === CODE BLOCKS === */
    .stCodeBlock {{
        background-color: {cor_fundo_secundario} !important;
    }}

    /* === TABS === */
    .stTabs [data-baseweb="tab-list"] button {{
        color: {cor_texto_secundario} !important;
    }}

    .stTabs [data-baseweb="tab-list"] button[aria-selected="true"] {{
        color: {cor_texto} !important;
        border-color: {cor_accent} !important;
    }}

    /* === FILE UPLOADER === */
    .stFileUploader > div {{
        border-color: {cor_borda} !important;
        background-color: {cor_card} !important;
    }}

    /* === SCROLLBAR === */
    ::-webkit-scrollbar {{
        width: 8px;
    }}

    ::-webkit-scrollbar-track {{
        background: {cor_fundo_secundario};
    }}

    ::-webkit-scrollbar-thumb {{
        background: {cor_borda};
        border-radius: 4px;
    }}

    </style>

    <script>
    try {{
        var meta = document.getElementById('pwa-theme-color');
        if (meta) {{
            meta.setAttribute('content', '{pwa_color}');
        }}
    }} catch(e) {{}}
    </script>
    """

    st.markdown(tema_css, unsafe_allow_html=True)


# ============================================================
# SIDEBAR — TOGGLE DE TEMA
# ============================================================

with st.sidebar:
    st.markdown("⚙️ **Configurações**")
    st.markdown("---")

    tema_atual = st.session_state.get("theme_mode", "dark")

    novo_tema = st.radio(
        "Tema da interface:",
        options=["dark", "light"],
        format_func=lambda x: "🌙 Escuro" if x == "dark" else "☀️ Claro",
        index=0 if tema_atual == "dark" else 1,
        key="toggle_tema"
    )

    if novo_tema != tema_atual:
        st.session_state.theme_mode = novo_tema
        st.rerun()

    st.markdown("---")
    st.caption("DocScripta AI v1.0")
    st.caption("Criado por: Eristionny 🤖")


# Aplica o tema CSS
aplicar_tema()


# ============================================================
# FUNÇÕES DE ARQUIVO
# ============================================================

def extrair_texto_arquivo(arquivo):

    if arquivo is None:
        return ""

    nome = arquivo.name.lower()

    try:

        if nome.endswith(".docx"):

            from docx import Document

            documento = Document(arquivo)

            partes = []

            for paragrafo in documento.paragraphs:

                texto = paragrafo.text.strip()

                if texto:
                    partes.append(texto)

            for tabela in documento.tables:

                for linha in tabela.rows:

                    celulas = []

                    for celula in linha.cells:

                        texto = celula.text.strip()

                        if texto:
                            celulas.append(texto)

                    if celulas:
                        partes.append(
                            " | ".join(celulas)
                        )

            return "\n".join(partes)

        if nome.endswith(".pdf"):

            from pypdf import PdfReader

            dados = arquivo.read()

            leitor = PdfReader(
                io.BytesIO(dados)
            )

            paginas = []

            for pagina in leitor.pages:

                texto = pagina.extract_text()

                if texto:
                    paginas.append(
                        texto.strip()
                    )

            return "\n".join(paginas)

        return ""

    except Exception as erro:

        raise Exception(
            f"Não foi possível ler o arquivo: {erro}"
        )


# ============================================================
# CRIAR DOCX
# ============================================================

def criar_docx_com_texto(
    texto,
    nome_arquivo="Documento_Revisado_DocScripta.docx"
):

    from docx import Document

    documento = Document()

    for linha in texto.split("\n"):

        if linha.strip():

            documento.add_paragraph(
                linha
            )

    buffer = io.BytesIO()

    documento.save(buffer)

    buffer.seek(0)

    return buffer, nome_arquivo


# ============================================================
# GEMINI PARA REVISÃO DE REDAÇÃO
# ============================================================

def gerar_texto_com_gemini(prompt):

    try:

        from sintetizador import (
            chamar_gemini_com_fallback
        )

        return chamar_gemini_com_fallback(
            prompt
        )

    except ImportError as erro:

        raise Exception(
            "A função chamar_gemini_com_fallback "
            "não foi encontrada no sintetizador.py."
        ) from erro


# ============================================================
# SUBSTITUIÇÃO SEGURA
# ============================================================

def aplicar_substituicao_segura(
    texto_documento,
    trecho_original,
    trecho_novo
):

    if not texto_documento:

        return {
            "sucesso": False,
            "motivo": "Documento vazio.",
            "texto": texto_documento
        }

    if not trecho_original:

        return {
            "sucesso": False,
            "motivo": "Trecho original vazio.",
            "texto": texto_documento
        }

    if not trecho_novo:

        return {
            "sucesso": False,
            "motivo": "Nova versão vazia.",
            "texto": texto_documento
        }

    ocorrencias = texto_documento.count(
        trecho_original
    )

    if ocorrencias == 0:

        return {
            "sucesso": False,
            "motivo": (
                "O trecho original não foi encontrado "
                "exatamente no documento atual."
            ),
            "texto": texto_documento
        }

    if ocorrencias > 1:

        return {
            "sucesso": False,
            "motivo": (
                f"O trecho aparece {ocorrencias} vezes "
                "no documento atual. "
                "A substituição automática foi bloqueada "
                "para evitar alteração incorreta."
            ),
            "texto": texto_documento
        }

    novo_texto = texto_documento.replace(
        trecho_original,
        trecho_novo,
        1
    )

    return {
        "sucesso": True,
        "motivo": "Trecho substituído com sucesso.",
        "texto": novo_texto
    }


# ============================================================
# FUNÇÕES VISUAIS DO DETECTOR DE IA
# ============================================================

def obter_icone_nivel(nivel):

    nivel = (nivel or "").upper()

    if nivel == "MUITO ALTO":
        return "🔴"

    if nivel == "ALTO":
        return "🟠"

    if nivel == "MODERADO":
        return "🟡"

    if nivel in ["BAIXO", "MUITO BAIXO"]:
        return "🟢"

    return "⚪"


def mostrar_resultado_ia(resultado):

    if not resultado:
        return

    indice = float(
        resultado.get(
            "indice_estimado",
            0
        )
    )

    classificacao = (
        resultado.get(
            "classificacao",
            "NÃO INFORMADA"
        )
    )

    nivel_alerta = (
        resultado.get(
            "nivel_alerta",
            ""
        )
    )

    st.subheader(
        "🤖 Resultado da análise"
    )

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric(
            "Índice estimado",
            f"{indice:.1f}%"
        )

    with col2:
        st.metric(
            "Classificação",
            classificacao
        )

    with col3:
        st.metric(
            "Nível de alerta",
            nivel_alerta
        )

    if classificacao in [
        "MUITO ALTO",
        "ALTO"
    ]:

        st.warning(
            "⚠️ Foram identificados sinais estilísticos "
            "compatíveis com geração ou assistência por IA."
        )

    elif classificacao == "MODERADO":

        st.info(
            "⚠️ Foram identificados alguns sinais "
            "estilísticos que merecem revisão."
        )

    elif classificacao in [
        "BAIXO",
        "MUITO BAIXO"
    ]:

        st.success(
            "✅ Não foram identificados sinais fortes "
            "no texto analisado."
        )

    else:

        st.info(
            "Não foi possível determinar uma classificação forte."
        )

    st.caption(
        "Esta análise é baseada em características "
        "estilísticas e não constitui prova definitiva "
        "de autoria por IA."
    )

    # ========================================================
    # CARACTERÍSTICAS
    # ========================================================

    caracteristicas = resultado.get(
        "caracteristicas",
        []
    )

    st.subheader(
        "🧠 Características observadas"
    )

    if caracteristicas:

        for caracteristica in caracteristicas:

            st.write(
                f"✓ {caracteristica}"
            )

    else:

        st.write(
            "Nenhuma característica forte foi identificada."
        )

    # ========================================================
    # TRECHOS RELEVANTES
    # ========================================================

    trechos = resultado.get(
        "trechos_relevantes",
        resultado.get(
            "trechos",
            []
        )
    )

    st.subheader(
        "🔍 Trechos mais relevantes"
    )

    if not trechos:

        st.info(
            "Nenhum trecho relevante foi identificado."
        )

    else:

        for numero, trecho in enumerate(
            trechos,
            start=1
        ):

            indice_trecho = float(
                trecho.get(
                    "indice",
                    0
                )
            )

            nivel = (
                trecho.get(
                    "nivel",
                    "NÃO INFORMADO"
                )
            )

            icone = obter_icone_nivel(
                nivel
            )

            titulo = (
                f"{icone} Trecho {numero} — "
                f"{indice_trecho:.1f}% — {nivel}"
            )

            with st.expander(
                titulo,
                expanded=(
                    numero <= 3
                )
            ):

                st.markdown(
                    "**Texto sinalizado:**"
                )

                st.info(
                    trecho.get(
                        "texto",
                        ""
                    )
                )

                motivos = trecho.get(
                    "motivos",
                    []
                )

                if motivos:

                    st.markdown(
                        "**Motivos observados:**"
                    )

                    for motivo in motivos:

                        st.write(
                            f"• {motivo}"
                        )

                if nivel in [
                    "ALTO",
                    "MUITO ALTO"
                ]:

                    st.warning(
                        "⚠️ Este trecho apresenta maior "
                        "concentração dos sinais analisados."
                    )

    # ========================================================
    # MUDANÇAS DE ESTILO
    # ========================================================

    mudancas = resultado.get(
        "mudancas_estilo",
        []
    )

    if mudancas:

        st.subheader(
            "⚠️ Mudanças de estilo"
        )

        for numero, mudanca in enumerate(
            mudancas,
            start=1
        ):

            with st.expander(
                f"Mudança de estilo {numero}"
            ):

                st.write(
                    f"**Distância estimada:** "
                    f"{mudanca.get('distancia', 0)}"
                )

                st.markdown(
                    "**Bloco anterior:**"
                )

                st.info(
                    mudanca.get(
                        "paragrafo_anterior",
                        ""
                    )
                )

                st.markdown(
                    "**Bloco atual:**"
                )

                st.info(
                    mudanca.get(
                        "paragrafo_atual",
                        ""
                    )
                )

                st.caption(
                    "Uma mudança de estilo não significa, "
                    "por si só, uso de IA."
                )

    # ========================================================
    # DISTRIBUIÇÃO
    # ========================================================

    distribuicao = resultado.get(
        "distribuicao_sinais"
    )

    if distribuicao:

        st.subheader(
            "📊 Distribuição dos sinais"
        )

        ordem = [
            "baixo",
            "moderado",
            "alto",
            "muito_alto"
        ]

        nomes = {
            "baixo": "Baixo",
            "moderado": "Moderado",
            "alto": "Alto",
            "muito_alto": "Muito alto"
        }

        for chave in ordem:

            if chave in distribuicao:

                st.write(
                    f"**{nomes[chave]}:** "
                    f"{distribuicao[chave]}%"
                )

    # ========================================================
    # CONCENTRAÇÃO
    # ========================================================

    concentracao = resultado.get(
        "concentracao_documental"
    )

    if concentracao:

        st.subheader(
            "📄 Concentração dos sinais"
        )

        st.write(
            concentracao
        )

    # ========================================================
    # ESTATÍSTICAS
    # ========================================================

    estatisticas = resultado.get(
        "estatisticas",
        {}
    )

    if estatisticas:

        st.subheader(
            "📈 Estatísticas do texto"
        )

        col1, col2, col3 = st.columns(3)

        with col1:

            if "palavras" in estatisticas:

                st.metric(
                    "Palavras",
                    estatisticas["palavras"]
                )

            if "frases" in estatisticas:

                st.metric(
                    "Frases",
                    estatisticas["frases"]
                )

            if "paragrafos" in estatisticas:

                st.metric(
                    "Parágrafos",
                    estatisticas["paragrafos"]
                )

        with col2:

            if "media_palavras_frase" in estatisticas:

                st.metric(
                    "Média por frase",
                    estatisticas[
                        "media_palavras_frase"
                    ]
                )

            if "diversidade_lexical" in estatisticas:

                st.metric(
                    "Diversidade lexical",
                    estatisticas[
                        "diversidade_lexical"
                    ]
                )

        with col3:

            if "uniformidade" in estatisticas:

                st.metric(
                    "Uniformidade",
                    estatisticas[
                        "uniformidade"
                    ]
                )

            if "densidade_conectivos" in estatisticas:

                st.metric(
                    "Conectivos",
                    estatisticas[
                        "densidade_conectivos"
                    ]
                )


# ============================================================
# ESTADO INICIAL
# ============================================================

VALORES_INICIAIS = {

    "resultado_texto": None,
    "resultado_arquivo": None,

    "text_key": 0,

    # ORIGINALIDADE
    "texto_documento_revisao": "",
    "resultado_originalidade": None,
    "relatorio_originalidade": None,
    "trecho_revisao": "",
    "fonte_revisao": "",
    "resultado_parafrase": "",
    "resultado_antes_depois": None,
    "historico_alteracoes": [],
    "arquivo_carregado_nome": "",

    # DETECTOR DE IA
    "texto_analise_ia": "",
    "resultado_ia": None,
    "texto_revisao_ia": "",
    "resultado_antes_depois_ia": None,
    "historico_ia": [],
    "trecho_revisao_ia": "",
    "indice_trecho_revisao_ia": None,

    # TEMA
    "theme_mode": "dark",
}


for chave, valor in VALORES_INICIAIS.items():

    if chave not in st.session_state:

        st.session_state[chave] = valor


# ============================================================
# ATUALIZA EDITOR ORIGINALIDADE
# ============================================================

def atualizar_documento_pelo_editor():

    if "editor_documento" in st.session_state:

        st.session_state.texto_documento_revisao = (
            st.session_state.editor_documento
        )


# ============================================================
# ATUALIZA EDITOR IA
# ============================================================

def atualizar_editor_ia():

    if "editor_ia" in st.session_state:

        st.session_state.texto_analise_ia = (
            st.session_state.editor_ia
        )


# ============================================================
# CABEÇALHO
# ============================================================

# ============================================================
# TONNYBOT + TÍTULO
# ============================================================

col_tonny, col_titulo = st.columns([1, 5])

with col_tonny:
    try:
        st.image(
            "tonnybot.png",
            width=140
        )
    except Exception:
        st.image(
            "https://raw.githubusercontent.com/eristionny/DocScripta-AI/main/tonnybot.png",
            width=140
        )

with col_titulo:
    st.title("📚 DocScripta AI")
    st.caption(
        "Assistente acadêmico para pesquisa, resolução, "
        "revisão, originalidade e análise de IA."
    )

# "Criado por" movido para sidebar


# ============================================================
# ABAS
# ============================================================

aba_principal, aba_originalidade, aba_ia, aba_reescrever = st.tabs(
    [
        "🏠 DocScripta AI",
        "🛡️ Originalidade Acadêmica",
        "🤖 Análise de IA",
        "✏️ Reescrever Texto"
    ]
)


# ============================================================
# ABA 1 — PRINCIPAL
# ============================================================

with aba_principal:

    st.subheader(
        "Pesquisa e Resolução Acadêmica"
    )

    st.write(
        "Digite uma pergunta, atividade, trabalho "
        "ou envie um PDF/DOCX."
    )

    # ========================================================
    # PESQUISA DIRETA
    # ========================================================

    st.markdown(
        "### 🔎 Pesquisa Direta"
    )

    chave_form = (
        f"form_texto_{st.session_state.text_key}"
    )

    chave_campo = (
        f"texto_pesquisa_{st.session_state.text_key}"
    )

    with st.form(
        key=chave_form
    ):

        prompt_texto = st.text_input(
            "Digite o tema, pergunta ou comando:",
            key=chave_campo,
            placeholder=(
                "Digite sua pergunta..."
            )
        )

        # ----------------------------------------------------
        # BOTÕES
        # ----------------------------------------------------

        col_pesquisar, col_limpar = st.columns(2)

        with col_pesquisar:

            btn_texto = st.form_submit_button(
                "🔎 Pesquisar",
                use_container_width=True
            )

        with col_limpar:

            btn_limpar_pesquisa = st.form_submit_button(
                "🧹 Limpar pesquisa",
                use_container_width=True
            )

    # --------------------------------------------------------
    # LIMPAR PESQUISA
    # --------------------------------------------------------

    if btn_limpar_pesquisa:

        st.session_state.text_key += 1

        st.session_state.resultado_texto = None

        st.session_state.resultado_arquivo = None

        st.rerun()

    # --------------------------------------------------------
    # EXECUTAR PESQUISA
    # --------------------------------------------------------

    if (
        btn_texto
        and
        prompt_texto.strip()
    ):

        with st.spinner(
            "Analisando pesquisa..."
        ):

            try:

                (
                    texto_gerado,
                    nome_arquivo,
                    tipo
                ) = gerar_sintese_academica(
                    prompt_texto
                )

                st.session_state.resultado_texto = (
                    texto_gerado
                )

                st.session_state.resultado_arquivo = (
                    nome_arquivo
                )

                # ------------------------------------------------
                # LIMPA A BARRA DE PESQUISA APÓS SUCESSO
                # ------------------------------------------------

                st.session_state.text_key += 1

                st.rerun()

            except Exception as erro:

                st.error(
                    f"Erro na pesquisa: {erro}"
                )

    st.divider()

    # ========================================================
    # ARQUIVO
    # ========================================================

    st.markdown(
        "### 📄 Envio de Documento"
    )

    with st.form(
        "form_arquivo"
    ):

        arquivo = st.file_uploader(
            "Escolha um PDF ou DOCX:",
            type=[
                "pdf",
                "docx"
            ],
            key="arquivo_principal"
        )

        btn_arquivo = st.form_submit_button(
            "📖 Analisar Arquivo"
        )

    if (
        btn_arquivo
        and
        arquivo
    ):

        with st.spinner(
            "Lendo e processando o documento..."
        ):

            try:

                (
                    texto_gerado,
                    nome_arquivo,
                    tipo
                ) = gerar_sintese_academica(
                    arquivo
                )

                st.session_state.resultado_texto = (
                    texto_gerado
                )

                st.session_state.resultado_arquivo = (
                    nome_arquivo
                )

                st.success(
                    "Documento analisado com sucesso."
                )

            except Exception as erro:

                st.error(
                    f"Erro ao processar o documento: {erro}"
                )

    # ========================================================
    # RESULTADO DA PESQUISA
    # ========================================================

    if st.session_state.resultado_texto:

        st.divider()

        st.subheader(
            "📝 Resposta Gerada"
        )

        st.markdown(
            st.session_state.resultado_texto
        )

        if st.session_state.resultado_arquivo:

            try:

                with open(
                    st.session_state.resultado_arquivo,
                    "rb"
                ) as arquivo_word:

                    st.download_button(
                        label="📄 Baixar documento Word",
                        data=arquivo_word,
                        file_name=(
                            st.session_state.resultado_arquivo
                        ),
                        mime=(
                            "application/vnd.openxmlformats-"
                            "officedocument.wordprocessingml.document"
                        ),
                        key="download_resposta_principal"
                    )

            except FileNotFoundError:

                st.warning(
                    "Arquivo Word não encontrado."
                )


# ============================================================
# ABA 2 — ORIGINALIDADE
# ============================================================

with aba_originalidade:

    st.header(
        "🛡️ Originalidade Acadêmica"
    )

    st.write(
        "Analise, revise e melhore seu trabalho antes da entrega."
    )

    st.info(
        "Similaridade é um indicador em relação às fontes "
        "analisadas e não representa, isoladamente, "
        "uma conclusão de plágio."
    )

    # ========================================================
    # ARQUIVO
    # ========================================================

    st.subheader(
        "1. Documento"
    )

    arquivo_originalidade = st.file_uploader(
        "Enviar PDF/DOCX:",
        type=[
            "pdf",
            "docx"
        ],
        key="arquivo_originalidade_novo"
    )

    if arquivo_originalidade:

        nome_atual = (
            arquivo_originalidade.name
        )

        if (
            st.session_state.arquivo_carregado_nome
            !=
            nome_atual
        ):

            try:

                texto_extraido = (
                    extrair_texto_arquivo(
                        arquivo_originalidade
                    )
                )

                st.session_state.texto_documento_revisao = (
                    texto_extraido
                )

                st.session_state.arquivo_carregado_nome = (
                    nome_atual
                )

                st.session_state.resultado_originalidade = None
                st.session_state.relatorio_originalidade = None
                st.session_state.trecho_revisao = ""
                st.session_state.fonte_revisao = ""
                st.session_state.resultado_parafrase = ""
                st.session_state.resultado_antes_depois = None
                st.session_state.historico_alteracoes = []

                st.success(
                    f"Arquivo carregado: {nome_atual}"
                )

                st.rerun()

            except Exception as erro:

                st.error(
                    f"Erro ao ler o arquivo: {erro}"
                )

    # ========================================================
    # EDITOR
    # ========================================================

    st.markdown(
        "### ✏️ Texto atual do documento"
    )

    texto_atual_editor = st.text_area(
        "Edite o documento diretamente aqui:",
        value=(
            st.session_state.texto_documento_revisao
        ),
        height=380,
        key="editor_documento",
        on_change=(
            atualizar_documento_pelo_editor
        ),
        placeholder=(
            "Cole ou edite aqui o conteúdo completo "
            "do seu trabalho."
        )
    )

    st.session_state.texto_documento_revisao = (
        texto_atual_editor
    )

    quantidade_caracteres = len(
        texto_atual_editor
    )

    quantidade_palavras = len(
        texto_atual_editor.split()
    )

    col_info1, col_info2 = st.columns(2)

    with col_info1:

        st.caption(
            f"Caracteres: {quantidade_caracteres}"
        )

    with col_info2:

        st.caption(
            f"Palavras: {quantidade_palavras}"
        )

    # ========================================================
    # CONFIGURAÇÃO
    # ========================================================

    st.subheader(
        "2. Configuração da análise"
    )

    consultar_pubmed = st.checkbox(
        "🔬 Pesquisar fontes no PubMed/NCBI",
        value=True,
        key="consultar_pubmed"
    )

    quantidade_fontes = st.slider(
        "Quantidade de fontes:",
        min_value=1,
        max_value=10,
        value=5,
        key="quantidade_fontes"
    )

    # ========================================================
    # ANALISAR ORIGINALIDADE
    # ========================================================

    if st.button(
        "🔎 ANALISAR ORIGINALIDADE",
        type="primary",
        use_container_width=True,
        key="btn_analisar_originalidade"
    ):

        texto_para_analisar = (
            st.session_state.texto_documento_revisao
        )

        if not texto_para_analisar.strip():

            st.warning(
                "Insira um texto antes da análise."
            )

        else:

            with st.spinner(
                "Pesquisando fontes e analisando o documento..."
            ):

                try:

                    resultado = (
                        executar_analise_originalidade(
                            texto_para_analisar,
                            consultar_pubmed=(
                                consultar_pubmed
                            ),
                            max_resultados_pubmed=(
                                quantidade_fontes
                            )
                        )
                    )

                    st.session_state.resultado_originalidade = (
                        resultado
                    )

                    st.session_state.relatorio_originalidade = (
                        gerar_relatorio_originalidade(
                            resultado
                        )
                    )

                    st.session_state.trecho_revisao = ""
                    st.session_state.fonte_revisao = ""
                    st.session_state.resultado_parafrase = ""
                    st.session_state.resultado_antes_depois = None

                    st.success(
                        "Nova análise concluída."
                    )

                except Exception as erro:

                    st.error(
                        f"Erro na análise: {erro}"
                    )

    # ========================================================
    # RESULTADO
    # ========================================================

    resultado = (
        st.session_state.resultado_originalidade
    )

    if resultado:

        st.divider()

        st.subheader(
            "3. Resultado da análise"
        )

        indice = resultado.get(
            "indice_maximo",
            0
        )

        fontes = resultado.get(
            "fontes_analisadas",
            0
        )

        trechos = len(
            resultado.get(
                "trechos",
                []
            )
        )

        citacoes = len(
            resultado.get(
                "citacoes",
                []
            )
        )

        c1, c2, c3, c4 = st.columns(4)

        with c1:
            st.metric(
                "Similaridade",
                f"{indice:.2f}%"
            )

        with c2:
            st.metric(
                "Fontes",
                fontes
            )

        with c3:
            st.metric(
                "Trechos",
                trechos
            )

        with c4:
            st.metric(
                "Citações",
                citacoes
            )

        st.write(
            f"**Status:** "
            f"{resultado.get('status', '')}"
        )

        st.write(
            f"**Classificação:** "
            f"{resultado.get('classificacao', '')}"
        )

        # ====================================================
        # FONTES
        # ====================================================

        st.subheader(
            "📚 Fontes"
        )

        for numero, fonte in enumerate(
            resultado.get(
                "fontes",
                []
            ),
            start=1
        ):

            titulo = fonte.get(
                "titulo",
                "Fonte sem título"
            )

            similaridade_fonte = fonte.get(
                "similaridade"
            )

            pmid = fonte.get(
                "pmid",
                ""
            )

            doi = fonte.get(
                "doi",
                ""
            )

            link = fonte.get(
                "link",
                ""
            )

            with st.expander(
                f"{numero}. {titulo}"
            ):

                if similaridade_fonte is not None:

                    st.write(
                        f"Similaridade: "
                        f"**{similaridade_fonte:.2f}%**"
                    )

                if pmid:

                    st.write(
                        f"PMID: **{pmid}**"
                    )

                if doi:

                    st.write(
                        f"DOI: **{doi}**"
                    )

                if link:

                    st.markdown(
                        f"[🔗 Abrir fonte]({link})"
                    )

                st.write(
                    "Texto disponível: "
                    +
                    (
                        "Sim"
                        if fonte.get(
                            "tem_texto",
                            False
                        )
                        else
                        "Não"
                    )
                )

        # ====================================================
        # TRECHOS DE ORIGINALIDADE
        # ====================================================

        st.subheader(
            "⚠️ Trechos para revisão"
        )

        trechos_resultado = resultado.get(
            "trechos",
            []
        )

        if not trechos_resultado:

            st.success(
                "Nenhum trecho longo com coincidência "
                "foi encontrado."
            )

        else:

            for numero, item in enumerate(
                trechos_resultado,
                start=1
            ):

                trecho = item.get(
                    "trecho",
                    ""
                )

                fonte = item.get(
                    "fonte",
                    "Fonte não identificada"
                )

                palavras = item.get(
                    "quantidade_palavras",
                    0
                )

                with st.expander(
                    f"⚠️ Trecho {numero}"
                ):

                    st.write(
                        f"**Fonte:** {fonte}"
                    )

                    st.write(
                        f"**Palavras consecutivas:** "
                        f"{palavras}"
                    )

                    st.info(
                        trecho
                    )

                    if st.button(
                        "✍️ Selecionar para revisão",
                        key=f"selecionar_trecho_{numero}"
                    ):

                        st.session_state.trecho_revisao = (
                            trecho
                        )

                        st.session_state.fonte_revisao = (
                            fonte
                        )

                        st.session_state.resultado_parafrase = ""
                        st.session_state.resultado_antes_depois = None

                        st.rerun()

        # ====================================================
        # CITAÇÕES
        # ====================================================

        st.subheader(
            "📚 Possíveis citações"
        )

        citacoes_resultado = resultado.get(
            "citacoes",
            []
        )

        if not citacoes_resultado:

            st.success(
                "Nenhuma frase foi sinalizada automaticamente."
            )

        else:

            for numero, frase in enumerate(
                citacoes_resultado,
                start=1
            ):

                with st.expander(
                    f"⚠️ Possível citação {numero}"
                ):

                    st.write(
                        frase
                    )

        # ====================================================
        # REVISÃO DE ORIGINALIDADE
        # ====================================================

        if st.session_state.trecho_revisao:

            st.divider()

            st.subheader(
                "✍️ Revisar trecho"
            )

            st.write(
                f"**Fonte:** "
                f"{st.session_state.fonte_revisao}"
            )

            st.info(
                st.session_state.trecho_revisao
            )

            modo = st.radio(
                "Escolha a forma de revisão:",
                [
                    "Reescrever academicamente",
                    "Simplificar a linguagem",
                    "Transformar em citação direta"
                ],
                index=0,
                key="modo_revisao_originalidade"
            )

            if st.button(
                "✍️ GERAR NOVA VERSÃO",
                type="primary",
                key="btn_nova_versao_originalidade"
            ):

                trecho = (
                    st.session_state.trecho_revisao
                )

                fonte = (
                    st.session_state.fonte_revisao
                )

                if modo == (
                    "Reescrever academicamente"
                ):

                    prompt = (
                        construir_prompt_parafrase(
                            trecho,
                            fonte
                        )
                    )

                elif modo == (
                    "Simplificar a linguagem"
                ):

                    prompt = f"""
Reescreva o trecho abaixo em português do Brasil
com linguagem acadêmica clara e mais simples.

Preserve:

- significado;
- fatos;
- números;
- resultados;
- sentido científico;
- necessidade de citação.

Não invente informação.

TRECHO:

{trecho}

FONTE:

{fonte}

ENTREGUE:

NOVA VERSÃO:
"""

                else:

                    prompt = f"""
Transforme o trecho abaixo em uma sugestão de
citação direta acadêmica.

Não altere as palavras originais.

Não invente conteúdo.

TRECHO:

{trecho}

FONTE:

{fonte}

ENTREGUE:

CITAÇÃO:
"texto"

REFERÊNCIA:
{fonte}
"""

                try:

                    with st.spinner(
                        "Gerando nova versão..."
                    ):

                        nova_versao = (
                            gerar_texto_com_gemini(
                                prompt
                            )
                        )

                    st.session_state.resultado_parafrase = (
                        nova_versao
                    )

                    st.success(
                        "Nova versão gerada."
                    )

                except Exception as erro:

                    st.error(
                        f"Erro ao gerar nova versão: {erro}"
                    )

            if st.session_state.resultado_parafrase:

                st.markdown(
                    "### 📝 Nova versão"
                )

                st.success(
                    st.session_state.resultado_parafrase
                )

                if st.button(
                    "🔄 COMPARAR ANTES E DEPOIS",
                    key="btn_comparar_originalidade"
                ):

                    try:

                        comparacao = (
                            comparar_antes_depois(
                                st.session_state.trecho_revisao,
                                st.session_state.resultado_parafrase,
                                st.session_state.fonte_revisao
                            )
                        )

                        st.session_state.resultado_antes_depois = (
                            comparacao
                        )

                    except Exception as erro:

                        st.error(
                            f"Erro na comparação: {erro}"
                        )

                comparacao = (
                    st.session_state.resultado_antes_depois
                )

                if comparacao:

                    a1, a2, a3 = st.columns(3)

                    with a1:

                        st.metric(
                            "Antes",
                            f"{comparacao['antes']:.2f}%"
                        )

                    with a2:

                        st.metric(
                            "Depois",
                            f"{comparacao['depois']:.2f}%"
                        )

                    with a3:

                        st.metric(
                            "Diferença",
                            f"{comparacao['reducao']:.2f} p.p."
                        )

                if st.button(
                    "✅ APLICAR ALTERAÇÃO AO DOCUMENTO",
                    type="primary",
                    key="btn_aplicar_originalidade"
                ):

                    substituicao = (
                        aplicar_substituicao_segura(
                            st.session_state.texto_documento_revisao,
                            st.session_state.trecho_revisao,
                            st.session_state.resultado_parafrase
                        )
                    )

                    if substituicao["sucesso"]:

                        st.session_state.historico_alteracoes.append(
                            st.session_state.texto_documento_revisao
                        )

                        st.session_state.texto_documento_revisao = (
                            substituicao["texto"]
                        )

                        st.session_state.trecho_revisao = ""
                        st.session_state.fonte_revisao = ""
                        st.session_state.resultado_parafrase = ""
                        st.session_state.resultado_antes_depois = None

                        st.success(
                            "Alteração aplicada ao documento."
                        )

                        st.rerun()

                    else:

                        st.error(
                            substituicao["motivo"]
                        )

                if st.button(
                    "❌ Descartar nova versão",
                    key="btn_descartar_originalidade"
                ):

                    st.session_state.resultado_parafrase = ""
                    st.session_state.resultado_antes_depois = None

                    st.rerun()

        # ====================================================
        # DOCUMENTO ATUAL
        # ====================================================

        st.divider()

        st.subheader(
            "📄 Documento atual"
        )

        texto_documento_final = (
            st.session_state.texto_documento_revisao
        )

        st.text_area(
            "Conteúdo atual:",
            value=texto_documento_final,
            height=350,
            disabled=True,
            key="visualizacao_documento_atual"
        )

        # ====================================================
        # COPIAR
        # ====================================================

        st.markdown(
            "### 📋 Copiar documento atual"
        )

        texto_js = (
            texto_documento_final
            .replace("\\", "\\\\")
            .replace("`", "\\`")
            .replace("${", "\\${")
        )

        codigo_copia = f"""
        <script>
        function copiarDocScripta() {{

            const texto = `{texto_js}`;

            navigator.clipboard.writeText(texto)
                .then(function() {{

                    document.getElementById(
                        "aviso-docscripta"
                    ).style.display = "block";

                    setTimeout(function() {{

                        document.getElementById(
                            "aviso-docscripta"
                        ).style.display = "none";

                    }}, 2500);

                }})
                .catch(function() {{

                    alert(
                        "Não foi possível copiar automaticamente."
                    );

                }});
        }}
        </script>

        <button
            onclick="copiarDocScripta()"
            style="
                width:100%;
                padding:12px;
                border:none;
                border-radius:8px;
                cursor:pointer;
                font-size:16px;
                font-weight:600;
            "
        >
            📋 Copiar documento atual
        </button>

        <div
            id="aviso-docscripta"
            style="
                display:none;
                margin-top:10px;
                padding:10px;
                border-radius:8px;
                text-align:center;
                background:#d1fae5;
                color:#065f46;
                font-weight:600;
            "
        >
            ✅ Documento copiado!
        </div>
        """

        st.components.v1.html(
            codigo_copia,
            height=85
        )

        # ====================================================
        # DESFAZER
        # ====================================================

        if st.session_state.historico_alteracoes:

            if st.button(
                "↩️ Desfazer última alteração",
                key="btn_desfazer_originalidade"
            ):

                texto_anterior = (
                    st.session_state.historico_alteracoes.pop()
                )

                st.session_state.texto_documento_revisao = (
                    texto_anterior
                )

                st.success(
                    "Última alteração desfeita."
                )

                st.rerun()

        # ====================================================
        # REANALISAR
        # ====================================================

        st.subheader(
            "🔄 Reanalisar documento atual"
        )

        if st.button(
            "🔄 NOVA ANÁLISE DO DOCUMENTO",
            use_container_width=True,
            key="btn_reanalisar_originalidade"
        ):

            texto_atual = (
                st.session_state.texto_documento_revisao
            )

            if not texto_atual.strip():

                st.warning(
                    "O documento está vazio."
                )

            else:

                with st.spinner(
                    "Reanalisando documento atual..."
                ):

                    try:

                        novo_resultado = (
                            executar_analise_originalidade(
                                texto_atual,
                                consultar_pubmed=(
                                    consultar_pubmed
                                ),
                                max_resultados_pubmed=(
                                    quantidade_fontes
                                )
                            )
                        )

                        st.session_state.resultado_originalidade = (
                            novo_resultado
                        )

                        st.session_state.relatorio_originalidade = (
                            gerar_relatorio_originalidade(
                                novo_resultado
                            )
                        )

                        st.success(
                            "Nova análise concluída usando "
                            "o documento atual."
                        )

                    except Exception as erro:

                        st.error(
                            f"Erro na nova análise: {erro}"
                        )

        # ====================================================
        # DOWNLOAD WORD
        # ====================================================

        st.subheader(
            "📥 Exportação"
        )

        if texto_documento_final.strip():

            buffer_word, nome_word = (
                criar_docx_com_texto(
                    texto_documento_final
                )
            )

            st.download_button(
                label=(
                    "📄 Baixar documento revisado (.docx)"
                ),
                data=buffer_word,
                file_name=nome_word,
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                key="download_doc_revisado"
            )

        # ====================================================
        # RELATÓRIO ORIGINALIDADE
        # ====================================================

        st.subheader(
            "📊 Relatório de Originalidade"
        )

        if st.session_state.relatorio_originalidade:

            with st.expander(
                "Visualizar relatório"
            ):

                st.markdown(
                    st.session_state.relatorio_originalidade
                )

            st.download_button(
                label="📄 Baixar relatório (.txt)",
                data=(
                    st.session_state.relatorio_originalidade
                    .encode("utf-8")
                ),
                file_name=(
                    "Relatorio_Originalidade_DocScripta.txt"
                ),
                mime="text/plain",
                key="download_relatorio_originalidade"
            )

        # ====================================================
        # LIMPAR
        # ====================================================

        st.divider()

        if st.button(
            "🗑️ Limpar análise",
            key="btn_limpar_originalidade"
        ):

            st.session_state.texto_documento_revisao = ""
            st.session_state.resultado_originalidade = None
            st.session_state.relatorio_originalidade = None
            st.session_state.trecho_revisao = ""
            st.session_state.fonte_revisao = ""
            st.session_state.resultado_parafrase = ""
            st.session_state.resultado_antes_depois = None
            st.session_state.historico_alteracoes = []
            st.session_state.arquivo_carregado_nome = ""

            st.rerun()


# ============================================================
# ABA 3 — ANÁLISE DE IA
# ============================================================

with aba_ia:

    st.header(
        "🤖 Análise de IA"
    )

    st.write(
        "Analise características estilísticas compatíveis "
        "com geração ou assistência por IA."
    )

    st.info(
        "A análise é probabilística e não constitui prova "
        "definitiva de autoria por IA."
    )

    # ========================================================
    # TEXTO
    # ========================================================

    st.subheader(
        "1. Texto para análise"
    )

    texto_ia_editor = st.text_area(
        "Digite ou cole o texto para análise:",
        value=(
            st.session_state.texto_analise_ia
        ),
        height=380,
        key="editor_ia",
        on_change=(
            atualizar_editor_ia
        ),
        placeholder=(
            "Cole aqui o texto que deseja analisar."
        )
    )

    st.session_state.texto_analise_ia = (
        texto_ia_editor
    )

    # ========================================================
    # ARQUIVO
    # ========================================================

    arquivo_ia = st.file_uploader(
        "Ou envie um PDF/DOCX:",
        type=[
            "pdf",
            "docx"
        ],
        key="arquivo_ia"
    )

    if arquivo_ia:

        nome_arquivo_ia = (
            arquivo_ia.name
        )

        if st.session_state.get(
            "arquivo_ia_carregado",
            ""
        ) != nome_arquivo_ia:

            try:

                texto_extraido_ia = (
                    extrair_texto_arquivo(
                        arquivo_ia
                    )
                )

                st.session_state.texto_analise_ia = (
                    texto_extraido_ia
                )

                st.session_state.arquivo_ia_carregado = (
                    nome_arquivo_ia
                )

                st.success(
                    f"Arquivo carregado: {nome_arquivo_ia}"
                )

            except Exception as erro:

                st.error(
                    f"Erro ao ler o arquivo: {erro}"
                )

    # ========================================================
    # CONTADORES
    # ========================================================

    palavras_ia = len(
        st.session_state.texto_analise_ia.split()
    )

    caracteres_ia = len(
        st.session_state.texto_analise_ia
    )

    c1, c2 = st.columns(2)

    with c1:

        st.caption(
            f"Palavras: {palavras_ia}"
        )

    with c2:

        st.caption(
            f"Caracteres: {caracteres_ia}"
        )

    # ========================================================
    # ANALISAR
    # ========================================================

    if st.button(
        "🔎 ANALISAR TEXTO",
        type="primary",
        use_container_width=True,
        key="btn_analisar_ia"
    ):

        texto_para_ia = (
            st.session_state.texto_analise_ia
        )

        if not texto_para_ia.strip():

            st.warning(
                "Digite ou envie um texto antes da análise."
            )

        else:

            with st.spinner(
                "Analisando características estilísticas..."
            ):

                try:

                    resultado_ia = (
                        analisar_texto_ia(
                            texto_para_ia
                        )
                    )

                    st.session_state.resultado_ia = (
                        resultado_ia
                    )

                    st.session_state.texto_revisao_ia = ""
                    st.session_state.resultado_antes_depois_ia = None

                    st.success(
                        "Análise concluída."
                    )

                except Exception as erro:

                    st.error(
                        f"Erro na análise de IA: {erro}"
                    )

    # ========================================================
    # RESULTADO
    # ========================================================

    resultado_ia = (
        st.session_state.resultado_ia
    )

    if resultado_ia:

        mostrar_resultado_ia(
            resultado_ia
        )

        # ====================================================
        # TRECHOS PARA REVISÃO
        # ====================================================

        trechos_ia = resultado_ia.get(
            "trechos_relevantes",
            resultado_ia.get(
                "trechos",
                []
            )
        )

        if trechos_ia:

            st.divider()

            st.subheader(
                "✍️ Revisão de redação"
            )

            indice_selecionado = st.session_state.get(
                "indice_trecho_revisao_ia"
            )

            if (
                indice_selecionado is not None
                and
                indice_selecionado < len(trechos_ia)
            ):

                trecho_selecionado = (
                    trechos_ia[
                        indice_selecionado
                    ]
                )

                st.info(
                    trecho_selecionado.get(
                        "texto",
                        ""
                    )
                )

                st.write(
                    f"**Índice:** "
                    f"{trecho_selecionado.get('indice', 0):.1f}%"
                )

                st.write(
                    f"**Classificação:** "
                    f"{trecho_selecionado.get('nivel', '')}"
                )

            else:

                st.caption(
                    "Selecione um trecho abaixo para "
                    "fazer uma revisão de redação."
                )

            # ------------------------------------------------
            # SELEÇÃO DE TRECHO
            # ------------------------------------------------

            for numero, trecho in enumerate(
                trechos_ia,
                start=1
            ):

                if st.button(
                    f"✍️ Melhorar redação — Trecho {numero}",
                    key=f"btn_melhorar_trecho_{numero}"
                ):

                    st.session_state.indice_trecho_revisao_ia = (
                        numero - 1
                    )

                    st.session_state.texto_revisao_ia = ""
                    st.session_state.resultado_antes_depois_ia = None

                    st.rerun()

            # ------------------------------------------------
            # GERAR REVISÃO
            # ------------------------------------------------

            indice_selecionado = (
                st.session_state.get(
                    "indice_trecho_revisao_ia"
                )
            )

            if (
                indice_selecionado is not None
                and
                indice_selecionado < len(trechos_ia)
            ):

                trecho_selecionado = (
                    trechos_ia[
                        indice_selecionado
                    ]
                )

                texto_trecho_ia = (
                    trecho_selecionado.get(
                        "texto",
                        ""
                    )
                )

                if st.button(
                    "✍️ GERAR NOVA VERSÃO",
                    type="primary",
                    key="btn_gerar_revisao_ia"
                ):

                    prompt_revisao_ia = f"""
Reescreva o trecho abaixo com linguagem acadêmica
mais clara, natural e bem organizada.

OBJETIVO:

Melhorar a redação do texto.

PRESERVE INTEGRALMENTE:

- significado;
- fatos;
- números;
- resultados;
- informações técnicas;
- referências;
- citações;
- sentido original.

NÃO:

- invente informações;
- altere fatos;
- remova referências;
- tente burlar detectores;
- tente enganar ferramentas de análise.

TRECHO ORIGINAL:

{texto_trecho_ia}

ENTREGUE SOMENTE A NOVA VERSÃO DO TEXTO.
"""

                    try:

                        with st.spinner(
                            "Gerando nova versão..."
                        ):

                            nova_versao_ia = (
                                gerar_texto_com_gemini(
                                    prompt_revisao_ia
                                )
                            )

                        st.session_state.texto_revisao_ia = (
                            nova_versao_ia
                        )

                        st.success(
                            "Nova versão gerada."
                        )

                    except Exception as erro:

                        st.error(
                            f"Erro ao revisar o trecho: {erro}"
                        )

                # --------------------------------------------
                # RESULTADO DA REVISÃO
                # --------------------------------------------

                if st.session_state.texto_revisao_ia:

                    st.markdown(
                        "### ANTES"
                    )

                    st.text_area(
                        "Trecho original:",
                        value=texto_trecho_ia,
                        height=180,
                        disabled=True,
                        key="trecho_original_ia"
                    )

                    st.markdown(
                        "### DEPOIS"
                    )

                    st.text_area(
                        "Nova versão:",
                        value=(
                            st.session_state.texto_revisao_ia
                        ),
                        height=180,
                        disabled=True,
                        key="trecho_novo_ia"
                    )

                    col_a, col_b = st.columns(2)

                    with col_a:

                        if st.button(
                            "✅ ACEITAR ALTERAÇÃO",
                            type="primary",
                            key="btn_aceitar_ia"
                        ):

                            documento_atual = (
                                st.session_state.texto_analise_ia
                            )

                            substituicao = (
                                aplicar_substituicao_segura(
                                    documento_atual,
                                    texto_trecho_ia,
                                    st.session_state.texto_revisao_ia
                                )
                            )

                            if substituicao["sucesso"]:

                                st.session_state.historico_ia.append(
                                    documento_atual
                                )

                                st.session_state.texto_analise_ia = (
                                    substituicao["texto"]
                                )

                                st.session_state.texto_revisao_ia = ""
                                st.session_state.resultado_antes_depois_ia = None
                                st.session_state.indice_trecho_revisao_ia = None

                                st.success(
                                    "Alteração aplicada ao documento."
                                )

                                st.rerun()

                            else:

                                st.error(
                                    substituicao["motivo"]
                                )

                    with col_b:

                        if st.button(
                            "❌ DESCARTAR",
                            key="btn_descartar_ia"
                        ):

                            st.session_state.texto_revisao_ia = ""
                            st.session_state.resultado_antes_depois_ia = None

                            st.rerun()

        # ====================================================
        # DOCUMENTO ATUAL
        # ====================================================

        st.divider()

        st.subheader(
            "📄 Documento atual"
        )

        st.text_area(
            "Texto atual:",
            value=(
                st.session_state.texto_analise_ia
            ),
            height=350,
            disabled=True,
            key="documento_atual_ia"
        )

        # ====================================================
        # REANALISAR
        # ====================================================

        if st.button(
            "🔄 REANALISAR DOCUMENTO",
            use_container_width=True,
            key="btn_reanalisar_ia"
        ):

            texto_atual_ia = (
                st.session_state.texto_analise_ia
            )

            if not texto_atual_ia.strip():

                st.warning(
                    "O documento está vazio."
                )

            else:

                with st.spinner(
                    "Reanalisando documento atualizado..."
                ):

                    try:

                        novo_resultado_ia = (
                            analisar_texto_ia(
                                texto_atual_ia
                            )
                        )

                        st.session_state.resultado_ia = (
                            novo_resultado_ia
                        )

                        st.session_state.texto_revisao_ia = ""
                        st.session_state.indice_trecho_revisao_ia = None

                        st.success(
                            "Nova análise concluída."
                        )

                    except Exception as erro:

                        st.error(
                            f"Erro na reanálise: {erro}"
                        )

        # ====================================================
        # DESFAZER
        # ====================================================

        if st.session_state.historico_ia:

            if st.button(
                "↩️ DESFAZER ÚLTIMA ALTERAÇÃO",
                key="btn_desfazer_ia"
            ):

                texto_anterior = (
                    st.session_state.historico_ia.pop()
                )

                st.session_state.texto_analise_ia = (
                    texto_anterior
                )

                st.session_state.texto_revisao_ia = ""
                st.session_state.indice_trecho_revisao_ia = None

                st.success(
                    "Última alteração desfeita."
                )

                st.rerun()

        # ====================================================
        # COPIAR
        # ====================================================

        st.markdown(
            "### 📋 Copiar texto atual"
        )

        texto_ia_js = (
            st.session_state.texto_analise_ia
            .replace("\\", "\\\\")
            .replace("`", "\\`")
            .replace("${", "\\${")
        )

        codigo_copia_ia = f"""
        <script>
        function copiarTextoIA() {{

            const texto = `{texto_ia_js}`;

            navigator.clipboard.writeText(texto)
                .then(function() {{

                    document.getElementById(
                        "aviso-copia-ia"
                    ).style.display = "block";

                    setTimeout(function() {{

                        document.getElementById(
                            "aviso-copia-ia"
                        ).style.display = "none";

                    }}, 2500);

                }})
                .catch(function() {{

                    alert(
                        "Não foi possível copiar automaticamente."
                    );

                }});
        }}
        </script>

        <button
            onclick="copiarTextoIA()"
            style="
                width:100%;
                padding:12px;
                border:none;
                border-radius:8px;
                cursor:pointer;
                font-size:16px;
                font-weight:600;
            "
        >
            📋 Copiar texto atual
        </button>

        <div
            id="aviso-copia-ia"
            style="
                display:none;
                margin-top:10px;
                padding:10px;
                border-radius:8px;
                text-align:center;
                background:#d1fae5;
                color:#065f46;
                font-weight:600;
            "
        >
            ✅ Texto copiado!
        </div>
        """

        st.components.v1.html(
            codigo_copia_ia,
            height=85
        )

        # ====================================================
        # DOWNLOAD WORD
        # ====================================================

        buffer_ia, nome_ia = (
            criar_docx_com_texto(
                st.session_state.texto_analise_ia,
                "Documento_Analisado_IA_DocScripta.docx"
            )
        )

        st.download_button(
            label="📄 Baixar documento revisado (.docx)",
            data=buffer_ia,
            file_name=nome_ia,
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            key="download_doc_ia"
        )

        # ====================================================
        # RELATÓRIO
        # ====================================================

        st.subheader(
            "📊 Relatório da Análise de IA"
        )

        try:

            relatorio_ia = (
                gerar_relatorio_ia(
                    resultado_ia
                )
            )

            with st.expander(
                "Visualizar relatório"
            ):

                st.markdown(
                    relatorio_ia
                )

            st.download_button(
                label="📄 Baixar relatório da análise",
                data=relatorio_ia.encode(
                    "utf-8"
                ),
                file_name=(
                    "Relatorio_Analise_IA_DocScripta.txt"
                ),
                mime="text/plain",
                key="download_relatorio_ia"
            )

        except Exception as erro:

            st.warning(
                f"Não foi possível gerar o relatório: {erro}"
            )

    # ========================================================
    # LIMPAR ANÁLISE DE IA
    # ========================================================

    st.divider()

    if st.button(
        "🗑️ LIMPAR ANÁLISE DE IA",
        key="btn_limpar_ia"
    ):

        st.session_state.texto_analise_ia = ""
        st.session_state.resultado_ia = None
        st.session_state.texto_revisao_ia = ""
        st.session_state.resultado_antes_depois_ia = None
        st.session_state.historico_ia = []
        st.session_state.trecho_revisao_ia = ""
        st.session_state.indice_trecho_revisao_ia = None

        if "arquivo_ia_carregado" in st.session_state:
            st.session_state.arquivo_ia_carregado = ""

        st.rerun()



# ============================================================
# ABA 4 — REESCREVER TEXTO
# ============================================================

with aba_reescrever:

    st.subheader("✏️ Reescrever Texto")

    st.write(
        "Reescreva seu texto mantendo o significado "
        "e as informações importantes."
    )

    texto_reescrever = st.text_area(
        "Cole aqui o texto que deseja reescrever:",
        height=300,
        key="texto_reescrever"
    )

    tom_reescrever = st.selectbox(
        "Escolha o tom:",
        [
            "academico",
            "claro",
            "conciso",
            "criativo",
            "profissional"
        ],
        format_func=lambda x: {
            "academico": "🎓 Acadêmico",
            "claro": "💡 Claro e simples",
            "conciso": "📌 Conciso e direto",
            "criativo": "✨ Criativo",
            "profissional": "💼 Profissional"
        }[x],
        key="tom_reescrever"
    )

    nivel_reescrever = st.selectbox(
        "Quanto deseja alterar o texto?",
        [
            "leve",
            "medio",
            "forte"
        ],
        format_func=lambda x: {
            "leve": "🟢 Leve — pequenas alterações",
            "medio": "🟡 Médio — reestruturação das frases",
            "forte": "🔴 Forte — reescrita completa"
        }[x],
        key="nivel_reescrever"
    )

    btn_reescrever = st.button(
        "✏️ REESCREVER TEXTO",
        type="primary",
        use_container_width=True
    )

    if btn_reescrever:

        if not texto_reescrever.strip():

            st.warning(
                "Digite ou cole um texto antes de reescrever."
            )

        else:

            with st.spinner(
                "Reescrevendo o texto..."
            ):

                try:

                    (
                        texto_resultado,
                        prompt_utilizado,
                        tipo_resultado
                    ) = reescrever_texto(
                        texto_reescrever,
                        tom=tom_reescrever,
                        nivel=nivel_reescrever
                    )

                    st.session_state[
                        "texto_reescrito_resultado"
                    ] = texto_resultado

                    st.success(
                        "Texto reescrito com sucesso!"
                    )

                except Exception as erro:

                    st.error(
                        f"Erro ao reescrever o texto: {erro}"
                    )

    if st.session_state.get(
        "texto_reescrito_resultado"
    ):

        st.markdown("---")

        st.markdown(
            "### 📄 Texto reescrito"
        )

        resultado_reescrito = st.session_state[
            "texto_reescrito_resultado"
        ]

        st.text_area(
            "Resultado:",
            value=resultado_reescrito,
            height=400,
            key="resultado_reescrito_visualizacao"
        )

        st.download_button(
            "⬇️ Baixar texto reescrito",
            data=resultado_reescrito,
            file_name="texto_reescrito.txt",
            mime="text/plain",
            use_container_width=True
        )
