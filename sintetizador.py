# -*- coding: utf-8 -*-

import os
import re
from typing import Optional, Callable, List, Dict, Any

from google import genai
from google.genai import types

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from pypdf import PdfReader

from buscador import buscar_pubmed


# ============================================================
# DOCSCRIPTA AI
# MOTOR ACADÊMICO
# ============================================================
#
# Gemini = cérebro
# DocScripta = coordenador acadêmico
# Pesquisa = ferramenta auxiliar
# V/F = validação específica
# Limpeza final = proteção contra saída intermediária
#
# ============================================================


# ============================================================
# 1. CONFIGURAÇÃO GEMINI
# ============================================================

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    raise Exception(
        "ERRO: A variável GEMINI_API_KEY não foi encontrada.\n\n"
        "No PowerShell do VS Code, configure sua chave com:\n\n"
        '$env:GEMINI_API_KEY="SUA_CHAVE"\n'
    )

client = genai.Client(
    api_key=GEMINI_API_KEY
)


# ============================================================
# 2. MODELOS
# ============================================================

MODELO_RAPIDO = "gemini-3.5-flash-lite"
MODELO_COMPLEXO = "gemini-3.6-flash"

MODELO_RAPIDO_CACHE = MODELO_RAPIDO


# ============================================================
# 3. LIMITES
# ============================================================

MAX_TOKENS_RESPOSTA = 2500
MAX_TOKENS_ATIVIDADE = 5000
MAX_TOKENS_RELATORIO = 6500
MAX_ARTIGOS_PUBMED = 3


# ============================================================
# 4. INSTRUÇÃO GLOBAL
# ============================================================

SYSTEM_INSTRUCTION_PT = """
Você é o DocScripta AI, um especialista acadêmico.

O Gemini é o cérebro principal do sistema.

Sua função é compreender o pedido, interpretar o contexto,
resolver a tarefa, verificar a resposta e entregar somente
o conteúdo final destinado ao usuário.

REGRAS GERAIS:

1. Responda em português do Brasil.
2. Seja claro, objetivo e organizado.
3. Preserve o contexto original do aluno.
4. Resolva integralmente a tarefa.
5. Respeite todos os critérios do enunciado.
6. Não ignore perguntas ou subitens.
7. Não invente referências.
8. Não invente autores.
9. Não invente DOI.
10. Não invente PMID.
11. Não invente links.
12. Não exponha raciocínio interno.
13. Não exponha pensamentos privados.
14. Não mostre comentários sobre o processo de geração.
15. Preserve a numeração original.
16. Use Markdown simples e consistente.

IMPORTANTE:

Não escreva frases como:

"Analisando a alternativa..."
"Considerando a alternativa..."
"Vou verificar..."
"Vou conferir..."
"Vou comparar..."
"Vou pesquisar..."
"Com adaptação para o negrito..."
"Autor correto"
"Título correto"
"Edição correta"

A resposta deve ser apresentada diretamente ao usuário.

QUESTÕES V/F:

Quando houver uma questão realmente de V/F:

1. Avalie cada afirmação.
2. Monte a sequência.
3. Compare com as alternativas.
4. Determine a letra correspondente.
5. Mostre a letra.
6. Mostre a sequência.
7. Mostre a alternativa correspondente.
8. Explique cada afirmação.

FORMATO:

# RESPOSTA CORRETA: X

**ALTERNATIVA X — V, V, F, V, V**

**SEQUÊNCIA CORRETA: V, V, F, V, V**

# ANÁLISE DAS AFIRMAÇÕES

**I — V (VERDADEIRO)**
Justificativa.

**II — V (VERDADEIRO)**
Justificativa.

Não entregue somente a letra.

Não entregue somente a sequência.

QUESTÕES NORMAIS:

Se houver A, B, C, D e E sem sequência V/F:

# RESPOSTA CORRETA: X

**X) Texto da alternativa**

### EXPLICAÇÃO

...

ATIVIDADES CONTEXTUALIZADAS:

- Preserve o contexto.
- Use os dados específicos.
- Relacione cada resposta ao caso.
- Não responda genericamente.
- Resolva cada questão separadamente.

RELATÓRIOS / TCC / TRABALHOS:

- Respeite a estrutura solicitada.
- Respeite ABNT quando exigida.
- Respeite quantidade e período das referências.
- Não invente fontes.
"""


# ============================================================
# 5. LIMPEZA BÁSICA DA RESPOSTA DO GEMINI
# ============================================================

def limpar_resposta_gemini(texto: str) -> str:
    """
    Remove tags/metadados de pensamento e alguns marcadores
    técnicos sem destruir o conteúdo acadêmico.
    """

    if not texto:
        return ""

    padroes = [
        r"<thought>.*?</thought>",
        r"<thinking>.*?</thinking>",
        r"<reasoning>.*?</reasoning>",
        r"<analysis>.*?</analysis>",
    ]

    for padrao in padroes:
        texto = re.sub(
            padrao,
            "",
            texto,
            flags=re.DOTALL | re.IGNORECASE
        )

    termos_tecnicos = [
        "internal reasoning:",
        "chain of thought:",
        "hidden reasoning:",
        "system metadata:",
    ]

    linhas = []

    for linha in texto.splitlines():

        linha_lower = linha.lower()

        if any(
            termo in linha_lower
            for termo in termos_tecnicos
        ):
            continue

        linhas.append(linha)

    return "\n".join(linhas).strip()


# ============================================================
# 6. LIMPEZA PROGRAMÁTICA DA RESPOSTA FINAL
# ============================================================

def limpar_resposta_final(texto: str) -> str:
    """
    Limpeza programática da saída do Gemini.

    A função remove comentários sobre o processo de geração,
    mas preserva o conteúdo acadêmico útil.
    """

    if not texto:
        return ""

    texto = texto.replace("\r\n", "\n")
    texto = texto.replace("\r", "\n")

    linhas_originais = texto.split("\n")
    linhas_finais = []

    padroes_linha_processual = [
        r"^\s*(?:considerando|analisando)\s+a\s+alternativa\s+[A-E]"
        r"(?:\s*[:\-].*)?$",

        r"^\s*(?:considerando|analisando)\s+a\s+alternativa\s+[A-E]"
        r"\s*$",

        r"^\s*(?:vou|vamos)\s+(?:analisar|verificar|conferir|comparar|pesquisar)"
        r".*$",

        r"^\s*(?:estou|estamos)\s+(?:analisando|verificando|comparando|pesquisando)"
        r".*$",

        r"^\s*(?:preparando|processando|organizando)\s+(?:a\s+)?resposta.*$",

        r"^\s*(?:com|para)\s+adaptação\s+(?:para|do)\s+.*$",

        r"^\s*(?:adaptação\s+para\s+o\s+negrito).*$",

        r"^\s*(?:formatação\s+(?:correta|aplicada)).*$",

        r"^\s*(?:vou\s+colocar|colocando)\s+.*\s+negrito.*$",

        r"^\s*(?:meu\s+raciocínio|meu\s+pensamento|processo\s+interno).*$",

        r"^\s*(?:etapa\s+[0-9]+)\s*[:\-].*$",

        r"^\s*(?:analyzing|considering)\s+(?:option|alternative)\s+[A-E]"
        r"(?:\s*[:\-].*)?$",

        r"^\s*(?:i\s+will|let's)\s+(?:analyze|verify|check|compare|search).*$",

        r"^\s*(?:checking|verifying|formatting|adapting).*$",
    ]

    padroes_parenteses_remover = [
        r"\s*\(\s*autor\s+correto\s*\)",
        r"\s*\(\s*título\s+correto\s*\)",
        r"\s*\(\s*edição\s+correta\s*\)",
        r"\s*\(\s*local\s+correto\s*\)",
        r"\s*\(\s*editora\s+correta\s*\)",
        r"\s*\(\s*data\s+correta\s*\)",
        r"\s*\(\s*author\s+correct\s*\)",
        r"\s*\(\s*correct\s+title\s*\)",
    ]

    dentro_parenteses_processuais = [
        "autor correto",
        "título correto",
        "edição correta",
        "local correto",
        "editora correta",
        "data correta",
        "adaptação para o negrito",
        "adaptação para negrito",
        "formatação correta",
    ]

    for linha in linhas_originais:

        linha_limpa = linha.strip()

        if not linha_limpa:
            linhas_finais.append("")
            continue

        linha_lower = linha_limpa.lower()

        # Remove linhas puramente processuais.
        remover = False

        for padrao in padroes_linha_processual:

            if re.match(
                padrao,
                linha_limpa,
                flags=re.IGNORECASE
            ):
                remover = True
                break

        if remover:
            continue

        # Remove comentários entre parênteses sem remover
        # o conteúdo principal da linha.
        linha_processada = linha_limpa

        for padrao in padroes_parenteses_remover:

            linha_processada = re.sub(
                padrao,
                "",
                linha_processada,
                flags=re.IGNORECASE
            )

        # Linhas que consistem apenas em checklist interno.
        linha_sem_marcacao = re.sub(
            r"^[\-\*\•]\s*",
            "",
            linha_processada
        ).strip()

        if (
            linha_sem_marcacao.lower()
            in
            dentro_parenteses_processuais
        ):
            continue

        # Remove comentário específico após dois pontos quando
        # ele for claramente somente uma observação sobre
        # formatação, mas preserva o conteúdo antes dele.
        linha_processada = re.sub(
            r"\s*\(com adaptação.*?\)\s*",
            "",
            linha_processada,
            flags=re.IGNORECASE
        )

        linha_processada = re.sub(
            r"\s*\(adaptação para o negrito.*?\)\s*",
            "",
            linha_processada,
            flags=re.IGNORECASE
        )

        if linha_processada.strip():
            linhas_finais.append(
                linha_processada.strip()
            )

    # Remove excesso de linhas vazias.
    resultado = "\n".join(
        linhas_finais
    )

    resultado = re.sub(
        r"\n{3,}",
        "\n\n",
        resultado
    )

    return resultado.strip()


# ============================================================
# 7. VALIDAÇÃO DA RESPOSTA FINAL
# ============================================================

def validar_resposta_final(texto: str) -> str:
    """
    Verifica se ainda existem comentários processuais
    óbvios na saída.
    """

    if not texto:
        return ""

    proibidos = [
        "analisando a alternativa",
        "considerando a alternativa",
        "com adaptação para o negrito",
        "adaptação para o negrito",
        "autor correto",
        "título correto",
        "edição correta",
        "vou verificar a alternativa",
        "vou analisar a alternativa",
        "meu raciocínio",
        "processo interno",
    ]

    resultado = texto

    for termo in proibidos:

        if termo in resultado.lower():

            resultado = limpar_resposta_final(
                resultado
            )

            break

    resultado = re.sub(
        r"\n{3,}",
        "\n\n",
        resultado
    )

    return resultado.strip()


# ============================================================
# 8. LEITURA DE ARQUIVOS
# ============================================================

def ler_arquivo(caminho_ou_arquivo):

    nome = getattr(
        caminho_ou_arquivo,
        "name",
        str(caminho_ou_arquivo)
    ).lower()

    try:

        if nome.endswith(".docx"):

            doc = Document(
                caminho_ou_arquivo
            )

            conteudo = [
                p.text.strip()
                for p in doc.paragraphs
                if p.text.strip()
            ]

            for tabela in doc.tables:

                for linha in tabela.rows:

                    celulas = [
                        c.text.strip()
                        for c in linha.cells
                        if c.text.strip()
                    ]

                    if celulas:

                        conteudo.append(
                            " | ".join(celulas)
                        )

            return "\n".join(
                conteudo
            )

        if nome.endswith(".pdf"):

            reader = PdfReader(
                caminho_ou_arquivo
            )

            paginas = []

            for pagina in reader.pages:

                texto = pagina.extract_text()

                if texto:
                    paginas.append(
                        texto.strip()
                    )

            return "\n".join(
                paginas
            )

        return str(
            caminho_ou_arquivo
        )

    except Exception as e:

        raise Exception(
            f"Erro ao ler o arquivo ({nome}): {e}"
        )


# ============================================================
# 9. DETECÇÃO V/F
# ============================================================

def detectar_questao_vf(texto: str) -> bool:

    texto_upper = texto.upper()

    termos_vf = [
        "VERDADEIRO OU FALSO",
        "VERDADEIRO E FALSO",
        "VERDADEIRO/FALSO",
        "V OU F",
        "V/F",
        "JULGUE AS AFIRMATIVAS",
        "JULGUE AS AFIRMAÇÕES",
        "JULGUE AS AFIRMACOES",
        "CLASSIFIQUE AS AFIRMATIVAS",
        "CLASSIFIQUE AS AFIRMAÇÕES",
        "CLASSIFIQUE AS AFIRMACOES",
    ]

    if any(
        termo in texto_upper
        for termo in termos_vf
    ):
        return True

    padroes_alternativas = [
        r"\b[A-E]\s*[\):.\-]\s*[VF](?:\s*[,;/|\-]\s*[VF])+",
        r"\b[A-E]\s*[:\-]\s*[VF](?:\s*[,;/|\-]\s*[VF])+",
    ]

    for padrao in padroes_alternativas:

        if re.search(
            padrao,
            texto,
            re.IGNORECASE
        ):
            return True

    return False


# ============================================================
# 10. EXTRAÇÃO ALTERNATIVAS V/F
# ============================================================

def extrair_alternativas_vf(
    texto: str
) -> Dict[str, List[str]]:

    alternativas = {}

    padroes = [
        re.compile(
            r"(?:^|\n)\s*"
            r"([A-Ea-e])"
            r"\s*[\):.\-]\s*"
            r"((?:[VFvTf]\s*[,;/|\-]?\s*){2,})"
            r"(?=\n|\r|$)",
            re.MULTILINE
        ),

        re.compile(
            r"\b([A-Ea-e])"
            r"\s*[:\-]\s*"
            r"([VFvTf]"
            r"(?:\s*[,;/|\-]\s*[VFvTf])+)"
        ),
    ]

    for padrao in padroes:

        encontrados = padrao.findall(
            texto
        )

        for letra, bloco in encontrados:

            letra = letra.upper()

            sequencia = re.findall(
                r"[VF]",
                bloco.upper()
            )

            if sequencia:
                alternativas[letra] = sequencia

    return alternativas


# ============================================================
# 11. FORMATAR ALTERNATIVAS V/F
# ============================================================

def formatar_alternativas_vf(
    alternativas: Dict[str, List[str]]
) -> str:

    if not alternativas:
        return ""

    linhas = []

    for letra in [
        "A",
        "B",
        "C",
        "D",
        "E"
    ]:

        if letra not in alternativas:
            continue

        linhas.append(
            f"{letra}: "
            f"{', '.join(alternativas[letra])}"
        )

    return "\n".join(
        linhas
    )


# ============================================================
# 12. COMPARAÇÃO DA SEQUÊNCIA
# ============================================================

def encontrar_letra_por_sequencia(
    sequencia: List[str],
    alternativas: Dict[str, List[str]]
) -> Optional[str]:

    sequencia = [
        item.upper()
        for item in sequencia
        if item.upper() in ["V", "F"]
    ]

    if not sequencia:
        return None

    for letra, alternativa in alternativas.items():

        alternativa_limpa = [
            item.upper()
            for item in alternativa
            if item.upper() in ["V", "F"]
        ]

        if alternativa_limpa == sequencia:
            return letra

    return None


# ============================================================
# 13. EXTRAIR SEQUÊNCIA
# ============================================================

def extrair_sequencia_resposta(
    resposta: str
) -> List[str]:

    padroes = [
        r"SEQUÊNCIA(?:\s+CORRETA)?\s*[:\-]\s*([VF,\s;/|\-]+)",
        r"SEQUENCIA(?:\s+CORRETA)?\s*[:\-]\s*([VF,\s;/|\-]+)",
    ]

    for padrao in padroes:

        match = re.search(
            padrao,
            resposta,
            re.IGNORECASE
        )

        if match:

            sequencia = re.findall(
                r"[VF]",
                match.group(1).upper()
            )

            if sequencia:
                return sequencia

    return []


# ============================================================
# 14. CORREÇÃO AUTOMÁTICA V/F
# ============================================================

def corrigir_resultado_vf(
    resposta: str,
    pergunta_original: str
) -> str:

    if not detectar_questao_vf(
        pergunta_original
    ):
        return resposta

    alternativas = (
        extrair_alternativas_vf(
            pergunta_original
        )
    )

    if not alternativas:
        return resposta

    sequencia = (
        extrair_sequencia_resposta(
            resposta
        )
    )

    if not sequencia:
        return resposta

    letra = (
        encontrar_letra_por_sequencia(
            sequencia,
            alternativas
        )
    )

    if not letra:
        return resposta

    alternativa_formatada = ", ".join(
        alternativas[letra]
    )

    resposta_limpa = re.sub(
        r"(?im)^\s*#?\s*RESPOSTA CORRETA:.*$",
        "",
        resposta
    )

    resposta_limpa = re.sub(
        r"(?im)^\s*\*\*ALTERNATIVA .*?\*\*\s*$",
        "",
        resposta_limpa
    )

    resposta_limpa = re.sub(
        r"(?im)^\s*\*\*SEQUÊNCIA.*?\*\*\s*$",
        "",
        resposta_limpa
    )

    resposta_limpa = re.sub(
        r"(?im)^\s*\*\*SEQUENCIA.*?\*\*\s*$",
        "",
        resposta_limpa
    )

    cabecalho = (
        f"# RESPOSTA CORRETA: {letra}\n\n"
        f"**ALTERNATIVA {letra} — "
        f"{alternativa_formatada}**\n\n"
        f"**SEQUÊNCIA CORRETA: "
        f"{alternativa_formatada}**"
    )

    resultado = (
        cabecalho
        + "\n\n"
        + resposta_limpa.strip()
    )

    return limpar_resposta_final(
        resultado
    )


# ============================================================
# 15. ANALISADOR DE CRITÉRIOS
# ============================================================

class AnalisadorCriterios:

    def __init__(
        self,
        texto: str
    ):

        self.texto = texto
        self.tipo = "questao"

        self.eh_relatorio = False
        self.eh_atividade = False
        self.eh_contextualizada = False
        self.eh_estudo_caso = False
        self.eh_verdadeiro_falso = False

        self.precisa_pubmed = False
        self.precisa_google = False

        self.exige_referencias = False
        self.proibe_pesquisa_externa = False

        self.numero_referencias = None
        self.periodo_referencias = None
        self.limite_linhas = None
        self.limite_palavras = None

        self.norma_referencia = None

        self.requisitos = []

    def analisar(self):

        texto = self.texto.lower()

        # ----------------------------------------------------
        # RELATÓRIO
        # ----------------------------------------------------

        if any(
            termo in texto
            for termo in [
                "relatório",
                "relatorio",
                "tcc",
                "monografia",
                "trabalho acadêmico",
                "trabalho academico",
                "resenha",
            ]
        ):

            self.tipo = "relatorio"
            self.eh_relatorio = True

        # ----------------------------------------------------
        # ESTUDO DE CASO
        # ----------------------------------------------------

        elif "estudo de caso" in texto:

            self.tipo = "estudo_de_caso"
            self.eh_estudo_caso = True
            self.eh_contextualizada = True
            self.eh_atividade = True

        # ----------------------------------------------------
        # CONTEXTUALIZADA
        # ----------------------------------------------------

        elif any(
            termo in texto
            for termo in [
                "atividade contextualizada",
                "situação-problema",
                "situacao-problema",
                "caso clínico",
                "caso clinico",
            ]
        ):

            self.tipo = "atividade_contextualizada"
            self.eh_contextualizada = True
            self.eh_atividade = True

        # ----------------------------------------------------
        # ATIVIDADE
        # ----------------------------------------------------

        elif any(
            termo in texto
            for termo in [
                "atividade",
                "questão",
                "questao",
                "questões",
                "questoes",
                "responda",
            ]
        ):

            self.tipo = "atividade"
            self.eh_atividade = True

        # ----------------------------------------------------
        # V/F
        # ----------------------------------------------------

        self.eh_verdadeiro_falso = (
            detectar_questao_vf(
                self.texto
            )
        )

        # ----------------------------------------------------
        # REFERÊNCIAS
        # ----------------------------------------------------

        if any(
            termo in texto
            for termo in [
                "referência",
                "referencia",
                "referências",
                "referencias",
                "bibliografia",
                "bibliográficas",
                "bibliograficas",
                "fonte",
                "fontes",
                "artigos científicos",
                "artigos cientificos",
            ]
        ):

            self.exige_referencias = True

        # ----------------------------------------------------
        # PROIBIR PESQUISA
        # ----------------------------------------------------

        if any(
            termo in texto
            for termo in [
                "somente com o material fornecido",
                "somente com o material enviado",
                "exclusivamente com o material",
                "não pesquise",
                "nao pesquise",
                "sem pesquisa externa",
                "não utilizar fontes externas",
                "nao utilizar fontes externas",
            ]
        ):

            self.proibe_pesquisa_externa = True

        # ----------------------------------------------------
        # PUBMED
        # ----------------------------------------------------

        if any(
            termo in texto
            for termo in [
                "pubmed",
                "artigo científico",
                "artigo cientifico",
                "artigos científicos",
                "artigos cientificos",
                "literatura científica",
                "literatura cientifica",
                "evidências científicas",
                "evidencias cientificas",
                "doi",
                "pmid",
            ]
        ):

            self.precisa_pubmed = True

        # ----------------------------------------------------
        # GOOGLE
        # ----------------------------------------------------

        if any(
            termo in texto
            for termo in [
                "atual",
                "atualmente",
                "hoje",
                "recentemente",
                "últimas informações",
                "ultimas informacoes",
                "lei atual",
                "legislação",
                "legislacao",
                "preço atual",
                "preco atual",
                "dados atuais",
            ]
        ):

            self.precisa_google = True

        # ----------------------------------------------------
        # QUANTIDADE REFERÊNCIAS
        # ----------------------------------------------------

        padrao_referencias = re.search(
            r"(\d+)\s*(?:referências|referencias|fontes|artigos)",
            texto
        )

        if padrao_referencias:

            try:
                self.numero_referencias = int(
                    padrao_referencias.group(1)
                )
            except Exception:
                pass

        # ----------------------------------------------------
        # PERÍODO
        # ----------------------------------------------------

        padrao_periodo = re.search(
            r"(?:últimos|ultimos)\s+(\d+)\s+anos",
            texto
        )

        if padrao_periodo:

            self.periodo_referencias = (
                padrao_periodo.group(1)
            )

        # ----------------------------------------------------
        # LINHAS
        # ----------------------------------------------------

        padrao_linhas = re.search(
            r"(?:máximo|maximo|até|ate)\s+(\d+)\s+linhas",
            texto
        )

        if padrao_linhas:

            self.limite_linhas = int(
                padrao_linhas.group(1)
            )

        # ----------------------------------------------------
        # PALAVRAS
        # ----------------------------------------------------

        padrao_palavras = re.search(
            r"(?:máximo|maximo|até|ate)\s+(\d+)\s+palavras",
            texto
        )

        if padrao_palavras:

            self.limite_palavras = int(
                padrao_palavras.group(1)
            )

        # ----------------------------------------------------
        # ABNT
        # ----------------------------------------------------

        if "abnt" in texto:

            self.norma_referencia = "ABNT"

        # ----------------------------------------------------
        # ESTRUTURA
        # ----------------------------------------------------

        if (
            "introdução" in texto
            or "introducao" in texto
        ):

            self.requisitos.append(
                "Incluir introdução"
            )

        if "objetivos" in texto:

            self.requisitos.append(
                "Incluir objetivos"
            )

        if (
            "conclusão" in texto
            or "conclusao" in texto
        ):

            self.requisitos.append(
                "Incluir conclusão"
            )

        if "desenvolvimento" in texto:

            self.requisitos.append(
                "Incluir desenvolvimento"
            )

        return self


# ============================================================
# 16. MOTOR
# ============================================================

class MotorResolucaoAtividade:

    def __init__(
        self,
        texto: str
    ):

        self.texto = texto.strip()

        self.criterios = (
            AnalisadorCriterios(
                self.texto
            ).analisar()
        )

    def escolher_modelo(self):

        c = self.criterios

        if c.eh_relatorio:
            return MODELO_COMPLEXO

        if c.eh_estudo_caso:
            return MODELO_COMPLEXO

        if c.eh_contextualizada:
            return MODELO_COMPLEXO

        if c.eh_verdadeiro_falso:
            return MODELO_COMPLEXO

        if c.precisa_pubmed:
            return MODELO_COMPLEXO

        if len(self.texto) > 12000:
            return MODELO_COMPLEXO

        return MODELO_RAPIDO


# ============================================================
# 17. GEMINI
# ============================================================

def chamar_gemini(
    prompt: str,
    modelo: Optional[str] = None,
    max_tokens: int = MAX_TOKENS_RESPOSTA,
    usar_google_search: bool = False,
    on_chunk: Optional[Callable] = None
) -> str:

    global MODELO_RAPIDO_CACHE

    if modelo is None:
        modelo = MODELO_RAPIDO_CACHE

    modelos = [modelo]

    if modelo != MODELO_COMPLEXO:
        modelos.append(
            MODELO_COMPLEXO
        )

    ultimo_erro = None

    for nome_modelo in modelos:

        try:

            nivel_pensamento = (
                "minimal"
                if nome_modelo == MODELO_RAPIDO
                else "medium"
            )

            config_kwargs = {
                "system_instruction":
                    SYSTEM_INSTRUCTION_PT,

                "max_output_tokens":
                    max_tokens,

                "thinking_config":
                    types.ThinkingConfig(
                        thinking_level=
                        nivel_pensamento
                    ),
            }

            if usar_google_search:

                config_kwargs["tools"] = [
                    types.Tool(
                        google_search=
                        types.GoogleSearch()
                    )
                ]

            config = types.GenerateContentConfig(
                **config_kwargs
            )

            stream = (
                client.models.generate_content_stream(
                    model=nome_modelo,
                    contents=prompt,
                    config=config
                )
            )

            partes = []

            for chunk in stream:

                if not chunk:
                    continue

                texto_chunk = getattr(
                    chunk,
                    "text",
                    None
                )

                if not texto_chunk:
                    continue

                partes.append(
                    texto_chunk
                )

                if on_chunk:
                    on_chunk(
                        texto_chunk
                    )

            resposta = "".join(
                partes
            ).strip()

            if not resposta:

                raise Exception(
                    "A API Gemini retornou resposta vazia."
                )

            MODELO_RAPIDO_CACHE = (
                nome_modelo
            )

            # =================================================
            # LIMPEZA OBRIGATÓRIA
            # =================================================

            resposta = limpar_resposta_gemini(
                resposta
            )

            resposta = limpar_resposta_final(
                resposta
            )

            resposta = validar_resposta_final(
                resposta
            )

            if not resposta:

                raise Exception(
                    "A resposta ficou vazia após a limpeza."
                )

            return resposta

        except Exception as e:

            ultimo_erro = e

            print(
                f"[Gemini] Falha no modelo "
                f"{nome_modelo}: {e}"
            )

    raise Exception(
        "Falha ao chamar a API Gemini.\n"
        f"Último erro: {ultimo_erro}"
    )


# ============================================================
# 18. FALLBACK COMPATÍVEL COM O APP.PY
# ============================================================

def chamar_gemini_com_fallback(
    prompt: str,
    modelo: Optional[str] = None,
    max_tokens: int = MAX_TOKENS_RESPOSTA,
    usar_google_search: bool = False
) -> str:

    return chamar_gemini(
        prompt=prompt,
        modelo=modelo,
        max_tokens=max_tokens,
        usar_google_search=usar_google_search
    )


# ============================================================
# 19. PUBMED
# ============================================================

def buscar_fontes_pubmed(
    texto: str,
    quantidade: int = MAX_ARTIGOS_PUBMED
):

    try:

        artigos = buscar_pubmed(
            texto[:300],
            max_resultados=quantidade
        )

        if not isinstance(
            artigos,
            list
        ):
            return []

        return artigos

    except Exception as e:

        print(
            f"[PubMed] Erro: {e}"
        )

        return []


# ============================================================
# 20. CONTEXTO DAS FONTES
# ============================================================

def montar_contexto_fontes(
    artigos: List[Dict[str, Any]]
) -> str:

    if not artigos:
        return ""

    resultado = (
        "\n\nFONTES ENCONTRADAS:\n"
    )

    for indice, artigo in enumerate(
        artigos,
        start=1
    ):

        titulo = artigo.get(
            "titulo",
            ""
        )

        citacao = artigo.get(
            "citacao_abnt",
            ""
        )

        pmid = artigo.get(
            "pmid",
            ""
        )

        doi = artigo.get(
            "doi",
            ""
        )

        link = (
            artigo.get("url")
            or artigo.get("link")
            or artigo.get("pmid_url")
            or ""
        )

        resultado += (
            f"\nFonte {indice}\n"
            f"Título: {titulo}\n"
            f"Citação: {citacao}\n"
            f"PMID: {pmid}\n"
            f"DOI: {doi}\n"
            f"Link: {link}\n"
        )

    return resultado


# ============================================================
# 21. REGRAS DOS CRITÉRIOS
# ============================================================

def montar_regras_criterios(
    criterios: AnalisadorCriterios
) -> str:

    regras = [
        f"Tipo: {criterios.tipo}"
    ]

    if criterios.eh_verdadeiro_falso:

        regras.append(
            "V/F: mostrar letra + sequência "
            "+ alternativa correspondente "
            "+ análise de cada afirmação."
        )

    if criterios.numero_referencias:

        regras.append(
            f"Referências exigidas: "
            f"{criterios.numero_referencias}"
        )

    if criterios.periodo_referencias:

        regras.append(
            f"Período: últimos "
            f"{criterios.periodo_referencias} anos"
        )

    if criterios.limite_linhas:

        regras.append(
            f"Limite: "
            f"{criterios.limite_linhas} linhas"
        )

    if criterios.limite_palavras:

        regras.append(
            f"Limite: "
            f"{criterios.limite_palavras} palavras"
        )

    if criterios.norma_referencia:

        regras.append(
            f"Norma: "
            f"{criterios.norma_referencia}"
        )

    if criterios.proibe_pesquisa_externa:

        regras.append(
            "Pesquisa externa PROIBIDA."
        )

    regras.extend(
        criterios.requisitos
    )

    return "\n".join(
        regras
    )


# ============================================================
# 22. PROMPT QUESTÃO
# ============================================================

def prompt_questao(
    texto,
    regras,
    fontes,
    eh_verdadeiro_falso=False,
    alternativas_detectadas=""
):

    bloco_vf = ""

    if eh_verdadeiro_falso:

        bloco_vf = f"""
ESTA É UMA QUESTÃO DE VERDADEIRO/FALSO.

ALTERNATIVAS DETECTADAS PELO SISTEMA:

{alternativas_detectadas}

Faça:

1. Avalie cada afirmação.
2. Determine V ou F.
3. Monte a sequência.
4. Compare com A-E.
5. Determine a letra.
6. Mostre letra e sequência.
7. Mostre a alternativa correspondente.
8. Explique cada afirmação.

NÃO escreva comentários sobre o processo.
NÃO escreva "analisando a alternativa".
NÃO escreva comentários sobre formatação.

FORMATO:

# RESPOSTA CORRETA: X

**ALTERNATIVA X — V, V, F, V, V**

**SEQUÊNCIA CORRETA: V, V, F, V, V**

# ANÁLISE DAS AFIRMAÇÕES

**I — V (VERDADEIRO)**
Justificativa.

**II — V (VERDADEIRO)**
Justificativa.

**III — F (FALSO)**
Justificativa.
"""

    return f"""
Resolva a questão acadêmica.

CRITÉRIOS:
{regras}

ENUNCIADO ORIGINAL:
--------------------------------------------------
{texto}
--------------------------------------------------

{fontes}

{bloco_vf}

REGRAS DA RESPOSTA:

1. Entregue diretamente a resposta final.
2. Não exponha análise do processo.
3. Não escreva "Analisando a alternativa..."
4. Não escreva "Considerando a alternativa..."
5. Não escreva comentários sobre negrito ou Markdown.
6. Se houver alternativas, informe a letra e o texto da alternativa.
7. Explique a resposta.
8. Se houver fontes verificadas, apresente-as.
9. Não invente referências.

FORMATO:

# RESPOSTA CORRETA: X

**X) Texto da alternativa**

# EXPLICAÇÃO

Explique de forma clara.

# REFERÊNCIAS

Somente fontes verificadas.
"""


# ============================================================
# 23. PROMPT ATIVIDADE
# ============================================================

def prompt_atividade(
    texto,
    regras,
    fontes,
    eh_verdadeiro_falso=False,
    alternativas_detectadas=""
):

    bloco_vf = ""

    if eh_verdadeiro_falso:

        bloco_vf = f"""
SE ALGUMA QUESTÃO FOR V/F:

Alternativas detectadas:

{alternativas_detectadas}

Mostrar:

**Resposta correta: X**

**Alternativa X — V, V, F, V, V**

**Sequência correta: V, V, F, V, V**

Depois analisar todas as afirmações.

Não alterar o formato das demais questões.

Não escrever comentários sobre o processo.
"""

    return f"""
Resolva integralmente a atividade.

CRITÉRIOS:
{regras}

ATIVIDADE ORIGINAL:
--------------------------------------------------
{texto}
--------------------------------------------------

{fontes}

{bloco_vf}

REGRAS:

1. Não ignore nenhuma questão.
2. Preserve a numeração.
3. Preserve o contexto.
4. Não responda somente com letras.
5. Explique todas as respostas.
6. Em V/F, mostrar letra + sequência.
7. Nas demais questões, use o formato adequado.
8. Não invente referências.
9. Respeite todos os limites.
10. Entregue somente a resposta final.
11. Não mostre comentários sobre como a resposta foi produzida.
"""


# ============================================================
# 24. PROMPT ESTUDO DE CASO
# ============================================================

def prompt_estudo_caso(
    texto,
    regras,
    fontes,
    alternativas_detectadas=""
):

    return f"""
Resolva integralmente o estudo de caso.

CRITÉRIOS:
{regras}

ESTUDO DE CASO:
--------------------------------------------------
{texto}
--------------------------------------------------

{fontes}

Preserve integralmente o contexto do caso.

Não responda genericamente.

Relacione cada resposta aos dados do caso.

Se alguma questão for V/F, utilize:

**Resposta correta: X**

**Alternativa X — V, V, F, V, V**

**Sequência correta: V, V, F, V, V**

e depois analise cada afirmação.

As demais questões devem manter o formato próprio
de estudo de caso.

Não exponha comentários de processo.

FORMATO:

# ESTUDO DE CASO — RESOLUÇÃO

## QUESTÃO 1

**RESPOSTA**

...

**ANÁLISE DO CASO**

...

**FUNDAMENTAÇÃO**

...

**CONCLUSÃO DA QUESTÃO**

...

## QUESTÃO 2

...

# SÍNTESE DO CASO

...

# REFERÊNCIAS

Somente fontes verificadas.
"""


# ============================================================
# 25. PROMPT RELATÓRIO
# ============================================================

def prompt_relatorio(
    texto,
    regras,
    fontes
):

    return f"""
Crie o relatório/trabalho acadêmico solicitado.

CRITÉRIOS:
{regras}

CONTEÚDO ORIGINAL:
--------------------------------------------------
{texto}
--------------------------------------------------

{fontes}

IMPORTANTE:

Esta é uma tarefa de relatório/trabalho acadêmico.

Não utilizar o formato de questão V/F,
exceto se houver uma atividade interna no relatório
que exija especificamente esse tipo de resposta.

Não mostrar comentários sobre o processo de geração.

FORMATO:

# 1. INTRODUÇÃO

# 2. OBJETIVOS

## 2.1 Objetivo geral

## 2.2 Objetivos específicos

# 3. DESENVOLVIMENTO

Utilize subtítulos adequados.

# 4. CONSIDERAÇÕES FINAIS

# 5. REFERÊNCIAS BIBLIOGRÁFICAS

REGRAS:

- Não inventar fontes.
- Não inventar autores.
- Não inventar DOI.
- Não inventar PMID.
- Respeitar os critérios do enunciado.
- Aplicar ABNT quando exigida.
"""


# ============================================================
# 26. PROMPT DE TRADUÇÃO
# ============================================================

def traduzir_tema_para_ingles(
    tema: str
) -> str:

    prompt = f"""
Traduza o tema acadêmico abaixo para inglês.

Retorne somente o termo ou expressão traduzida,
sem explicações.

Tema:
{tema}
"""

    try:

        return chamar_gemini(
            prompt,
            modelo=MODELO_RAPIDO,
            max_tokens=200,
            usar_google_search=False
        ).strip()

    except Exception:

        return tema


# ============================================================
# 27. FORMATADOR WORD
# ============================================================

def formatar_documento_com_destaque(
    texto_gerado,
    titulo_doc,
    tipo="resposta"
):

    doc = Document()

    # --------------------------------------------------------
    # MARGENS
    # --------------------------------------------------------

    for section in doc.sections:

        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --------------------------------------------------------
    # FONTE
    # --------------------------------------------------------

    style = doc.styles["Normal"]

    style.font.name = "Calibri"
    style.font.size = Pt(11)

    style.font.color.rgb = RGBColor(
        0x22,
        0x22,
        0x22
    )

    # ========================================================
    # CAPA DE RELATÓRIO
    # ========================================================

    if tipo == "relatorio":

        p = doc.add_paragraph()

        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        r = p.add_run(
            "INSTITUIÇÃO DE ENSINO SUPERIOR\n"
            "CURSO DE GRADUAÇÃO\n"
        )

        r.bold = True
        r.font.size = Pt(12)

        r.font.color.rgb = RGBColor(
            0x1B,
            0x36,
            0x5D
        )

        p = doc.add_paragraph()

        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        r = p.add_run(
            f"\n\n{titulo_doc.upper()}\n\n"
        )

        r.bold = True
        r.font.size = Pt(16)

        r.font.color.rgb = RGBColor(
            0x1B,
            0x36,
            0x5D
        )

        p = doc.add_paragraph()

        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        r = p.add_run(
            "Aluno(a): Eristionny Alves Batista\n\n"
        )

        r.font.size = Pt(12)

        doc.add_page_break()

    # ========================================================
    # CABEÇALHO
    # ========================================================

    p = doc.add_paragraph()

    r = p.add_run(
        f"DocScripta AI — {titulo_doc}"
    )

    r.bold = True
    r.font.size = Pt(18)

    r.font.color.rgb = RGBColor(
        0x1B,
        0x36,
        0x5D
    )

    doc.add_paragraph()

    # ========================================================
    # PROCESSAMENTO
    # ========================================================

    for linha in texto_gerado.split("\n"):

        linha_limpa = linha.strip()

        if not linha_limpa:
            continue

        # ----------------------------------------------------
        # TÍTULOS
        # ----------------------------------------------------

        if linha_limpa.startswith("#"):

            titulo = (
                linha_limpa
                .lstrip("#")
                .strip()
            )

            p = doc.add_paragraph()

            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)

            r = p.add_run(
                titulo
            )

            r.bold = True
            r.font.size = Pt(14)

            r.font.color.rgb = RGBColor(
                0x1B,
                0x36,
                0x5D
            )

            continue

        # ----------------------------------------------------
        # RESPOSTA DESTACADA
        # ----------------------------------------------------

        eh_resposta_principal = any(
            linha_limpa.upper().startswith(chave)
            for chave in [
                "RESPOSTA CORRETA",
                "RESPOSTA DIRETA",
                "ALTERNATIVA CORRETA",
                "GABARITO",
                "ALTERNATIVA A —",
                "ALTERNATIVA B —",
                "ALTERNATIVA C —",
                "ALTERNATIVA D —",
                "ALTERNATIVA E —",
                "SEQUÊNCIA CORRETA",
                "SEQUENCIA CORRETA",
            ]
        )

        if eh_resposta_principal:

            table = doc.add_table(
                rows=1,
                cols=1
            )

            table.alignment = (
                WD_TABLE_ALIGNMENT.CENTER
            )

            cell = table.cell(
                0,
                0
            )

            shading = parse_xml(
                f'<w:shd {nsdecls("w")} '
                f'w:fill="EAF2F8"/>'
            )

            cell._tc.get_or_add_tcPr().append(
                shading
            )

            border = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:left w:val="single" '
                f'w:sz="36" '
                f'w:space="0" '
                f'w:color="1B365D"/>'
                f'</w:tcBorders>'
            )

            cell._tc.get_or_add_tcPr().append(
                border
            )

            p = cell.paragraphs[0]

            r = p.add_run(
                linha_limpa
            )

            r.bold = True
            r.font.size = Pt(16)

            r.font.color.rgb = RGBColor(
                0x1B,
                0x36,
                0x5D
            )

            doc.add_paragraph()

            continue

        # ----------------------------------------------------
        # QUESTÃO
        # ----------------------------------------------------

        if re.match(
            r"^(QUESTÃO|QUESTAO)\s+\d+",
            linha_limpa,
            re.IGNORECASE
        ):

            p = doc.add_paragraph()

            p.paragraph_format.space_before = Pt(15)

            r = p.add_run(
                linha_limpa
            )

            r.bold = True
            r.font.size = Pt(14)

            r.font.color.rgb = RGBColor(
                0x1B,
                0x36,
                0x5D
            )

            continue

        # ----------------------------------------------------
        # TEXTO NORMAL
        # ----------------------------------------------------

        p = doc.add_paragraph()

        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15

        partes = linha_limpa.split("**")

        for indice, parte in enumerate(
            partes
        ):

            if not parte:
                continue

            r = p.add_run(
                parte
            )

            if indice % 2 == 1:

                r.bold = True

                r.font.color.rgb = RGBColor(
                    0x1B,
                    0x36,
                    0x5D
                )

    # ========================================================
    # NOME DO ARQUIVO
    # ========================================================

    if tipo == "relatorio":

        nome_arquivo = (
            "Documento_Academico_DocScripta.docx"
        )

    elif tipo in [
        "estudo_de_caso",
        "atividade_contextualizada"
    ]:

        nome_arquivo = (
            "Atividade_Contextualizada_DocScripta.docx"
        )

    else:

        nome_arquivo = (
            "Atividade_Resolvida_DocScripta.docx"
        )

    doc.save(
        nome_arquivo
    )

    return nome_arquivo


# ============================================================
# 28. FUNÇÃO PRINCIPAL
# ============================================================

def gerar_sintese_academica(
    prompt_ou_arquivo,
    on_chunk: Optional[Callable] = None
):

    # ========================================================
    # LEITURA
    # ========================================================

    texto = ler_arquivo(
        prompt_ou_arquivo
    ).strip()

    if not texto:

        raise Exception(
            "O conteúdo enviado está vazio."
        )

    # ========================================================
    # MOTOR
    # ========================================================

    motor = MotorResolucaoAtividade(
        texto
    )

    criterios = motor.criterios

    modelo = (
        motor.escolher_modelo()
    )

    print(
        f"[DocScripta] Tipo: {criterios.tipo}"
    )

    print(
        f"[DocScripta] V/F: "
        f"{criterios.eh_verdadeiro_falso}"
    )

    print(
        f"[DocScripta] Modelo: {modelo}"
    )

    # ========================================================
    # ALTERNATIVAS V/F
    # ========================================================

    alternativas_vf = {}
    texto_alternativas = ""

    if criterios.eh_verdadeiro_falso:

        alternativas_vf = (
            extrair_alternativas_vf(
                texto
            )
        )

        texto_alternativas = (
            formatar_alternativas_vf(
                alternativas_vf
            )
        )

        if texto_alternativas:

            print(
                "[DocScripta] Alternativas V/F:"
            )

            print(
                texto_alternativas
            )

    # ========================================================
    # FONTES
    # ========================================================

    artigos = []

    if (
        criterios.precisa_pubmed
        and
        not criterios.proibe_pesquisa_externa
    ):

        print(
            "[DocScripta] Pesquisando PubMed..."
        )

        quantidade = (
            criterios.numero_referencias
            or MAX_ARTIGOS_PUBMED
        )

        quantidade = min(
            quantidade,
            10
        )

        artigos = buscar_fontes_pubmed(
            texto,
            quantidade
        )

    contexto_fontes = (
        montar_contexto_fontes(
            artigos
        )
    )

    # ========================================================
    # REGRAS
    # ========================================================

    regras = montar_regras_criterios(
        criterios
    )

    # ========================================================
    # RELATÓRIO
    # ========================================================

    if criterios.eh_relatorio:

        prompt = prompt_relatorio(
            texto,
            regras,
            contexto_fontes
        )

        resposta = chamar_gemini(
            prompt,
            modelo=MODELO_COMPLEXO,
            max_tokens=MAX_TOKENS_RELATORIO,
            usar_google_search=(
                criterios.precisa_google
                and
                not criterios.proibe_pesquisa_externa
            ),
            on_chunk=on_chunk
        )

        titulo = (
            texto[:80]
            .replace("\n", " ")
        )

        arquivo = (
            formatar_documento_com_destaque(
                resposta,
                titulo,
                tipo="relatorio"
            )
        )

        return (
            resposta,
            arquivo,
            "relatorio"
        )

    # ========================================================
    # ESTUDO DE CASO / CONTEXTUALIZADA
    # ========================================================

    if criterios.tipo in [
        "estudo_de_caso",
        "atividade_contextualizada"
    ]:

        prompt = prompt_estudo_caso(
            texto,
            regras,
            contexto_fontes,
            texto_alternativas
        )

        resposta = chamar_gemini(
            prompt,
            modelo=MODELO_COMPLEXO,
            max_tokens=MAX_TOKENS_ATIVIDADE,
            usar_google_search=(
                criterios.precisa_google
                and
                not criterios.proibe_pesquisa_externa
            ),
            on_chunk=on_chunk
        )

        if criterios.eh_verdadeiro_falso:

            resposta = (
                corrigir_resultado_vf(
                    resposta,
                    texto
                )
            )

        arquivo = (
            formatar_documento_com_destaque(
                resposta,
                "Atividade Contextualizada",
                tipo=criterios.tipo
            )
        )

        return (
            resposta,
            arquivo,
            criterios.tipo
        )

    # ========================================================
    # ATIVIDADE
    # ========================================================

    if criterios.eh_atividade:

        prompt = prompt_atividade(
            texto,
            regras,
            contexto_fontes,
            criterios.eh_verdadeiro_falso,
            texto_alternativas
        )

        resposta = chamar_gemini(
            prompt,
            modelo=modelo,
            max_tokens=MAX_TOKENS_ATIVIDADE,
            usar_google_search=(
                criterios.precisa_google
                and
                not criterios.proibe_pesquisa_externa
            ),
            on_chunk=on_chunk
        )

        if criterios.eh_verdadeiro_falso:

            resposta = (
                corrigir_resultado_vf(
                    resposta,
                    texto
                )
            )

        arquivo = (
            formatar_documento_com_destaque(
                resposta,
                "Atividade Acadêmica Resolvida",
                tipo="atividade"
            )
        )

        return (
            resposta,
            arquivo,
            "atividade"
        )

    # ========================================================
    # QUESTÃO SIMPLES
    # ========================================================

    prompt = prompt_questao(
        texto,
        regras,
        contexto_fontes,
        criterios.eh_verdadeiro_falso,
        texto_alternativas
    )

    resposta = chamar_gemini(
        prompt,
        modelo=modelo,
        max_tokens=MAX_TOKENS_RESPOSTA,
        usar_google_search=(
            criterios.precisa_google
            and
            not criterios.proibe_pesquisa_externa
        ),
        on_chunk=on_chunk
    )

    if criterios.eh_verdadeiro_falso:

        resposta = (
            corrigir_resultado_vf(
                resposta,
                texto
            )
        )

    # Limpeza final após qualquer correção V/F.
    resposta = limpar_resposta_final(
        resposta
    )

    resposta = validar_resposta_final(
        resposta
    )

    arquivo = (
        formatar_documento_com_destaque(
            resposta,
            "Resposta Acadêmica",
            tipo="resposta"
        )
    )

    return (
        resposta,
        arquivo,
        "resposta"
    )